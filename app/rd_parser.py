"""
R&D Tax Credit Analysis Parser

Handles parsing of Excel, PDF, and DOCX files for R&D tax credit analysis.
Uses Gemini for document understanding and four-part test evaluation.
"""

import os
import io
import json
import uuid
import logging
from typing import List, Dict, Any, Optional
from enum import Enum
from dataclasses import dataclass, asdict
from datetime import datetime

import pandas as pd
from pydantic import BaseModel

# Try to import Gemini (use legacy SDK)
try:
    import google.generativeai as genai
    GEMINI_AVAILABLE = True
except ImportError:
    GEMINI_AVAILABLE = False
    genai = None

try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

logger = logging.getLogger(__name__)


# =============================================================================
# ENUMS AND DATA MODELS
# =============================================================================

class TestStatus(str, Enum):
    PASS = "pass"
    FAIL = "fail"
    NEEDS_REVIEW = "needs_review"
    MISSING_DATA = "missing_data"


class FourPartTestResult(BaseModel):
    """Results of the four-part test evaluation"""
    permitted_purpose: TestStatus = TestStatus.MISSING_DATA
    permitted_purpose_reasoning: str = ""
    elimination_uncertainty: TestStatus = TestStatus.MISSING_DATA
    elimination_uncertainty_reasoning: str = ""
    process_experimentation: TestStatus = TestStatus.MISSING_DATA
    process_experimentation_reasoning: str = ""
    technological_nature: TestStatus = TestStatus.MISSING_DATA
    technological_nature_reasoning: str = ""
    
    @property
    def pass_count(self) -> int:
        tests = [self.permitted_purpose, self.elimination_uncertainty, 
                 self.process_experimentation, self.technological_nature]
        return sum(1 for t in tests if t == TestStatus.PASS)
    
    @property
    def total_tests(self) -> int:
        return 4


class RDProject(BaseModel):
    """A project being evaluated for R&D tax credit qualification"""
    project_id: str
    project_name: str
    category: Optional[str] = None
    description: Optional[str] = None
    budget: Optional[float] = None
    four_part_test: FourPartTestResult = FourPartTestResult()
    confidence_score: float = 0.0
    missing_info: List[str] = []
    ai_summary: str = ""
    qualified: bool = False


class RDEmployee(BaseModel):
    """Employee data for R&D analysis"""
    employee_id: str
    name: str
    job_title: Optional[str] = None
    department: Optional[str] = None
    location: Optional[str] = None
    w2_wages: float = 0.0
    qre_wage_base: float = 0.0
    rd_allocation_percent: float = 0.0
    stock_compensation: float = 0.0
    severance: float = 0.0


class RDVendor(BaseModel):
    """Vendor data for contract research"""
    vendor_id: str
    vendor_name: str
    risk_bearer: str = ""
    ip_rights: str = ""
    country: str = ""
    qualified: bool = False


class RDExpense(BaseModel):
    """Expense/transaction data"""
    transaction_id: str
    vendor_id: Optional[str] = None
    description: str
    amount: float
    qre_amount: float = 0.0
    qualified: bool = False
    category: str = ""  # supplies, contract_research, wages


class GapItem(BaseModel):
    """Information gap that needs user attention"""
    gap_id: str
    category: str  # project, employee, vendor, documentation
    item_id: str
    item_name: str
    gap_type: str  # missing_data, needs_clarification, needs_documentation
    description: str
    required_info: List[str] = []
    priority: str = "medium"  # high, medium, low


class RDAnalysisSession(BaseModel):
    """Complete R&D analysis session"""
    session_id: str
    created_at: str
    company_name: str = ""
    industry: str = ""
    tax_year: int = 2024
    
    # Parsed data
    projects: List[RDProject] = []
    employees: List[RDEmployee] = []
    vendors: List[RDVendor] = []
    expenses: List[RDExpense] = []
    
    # Analysis results
    gaps: List[GapItem] = []
    total_qre: float = 0.0
    wage_qre: float = 0.0
    supply_qre: float = 0.0
    contract_qre: float = 0.0
    
    # Summary statistics
    total_employees: int = 0
    rd_employees: int = 0
    qualified_projects: int = 0
    
    # Status
    parsing_complete: bool = False
    analysis_complete: bool = False
    errors: List[str] = []


# =============================================================================
# FILE PARSERS
# =============================================================================

def parse_excel_file(file_content: bytes, filename: str) -> Dict[str, pd.DataFrame]:
    """Parse Excel file into dictionary of DataFrames by sheet name"""
    try:
        buffer = io.BytesIO(file_content)
        xl = pd.ExcelFile(buffer)
        
        sheets = {}
        for sheet_name in xl.sheet_names:
            try:
                df = pd.read_excel(xl, sheet_name=sheet_name)
                # Clean column names
                df.columns = [str(c).strip() for c in df.columns]
                sheets[sheet_name] = df
                logger.info(f"Parsed sheet '{sheet_name}' with {len(df)} rows")
            except Exception as e:
                logger.warning(f"Could not parse sheet '{sheet_name}': {e}")
        
        return sheets
    except Exception as e:
        logger.error(f"Error parsing Excel file {filename}: {e}")
        raise ValueError(f"Could not parse Excel file: {e}")


def parse_pdf_file(file_content: bytes, filename: str) -> str:
    """Extract text from PDF file"""
    if not PDF_AVAILABLE:
        raise ValueError("PDF parsing not available. Install PyPDF2.")
    
    try:
        buffer = io.BytesIO(file_content)
        reader = PyPDF2.PdfReader(buffer)
        
        text_parts = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)
        
        full_text = "\n\n".join(text_parts)
        logger.info(f"Extracted {len(full_text)} characters from PDF {filename}")
        return full_text
    except Exception as e:
        logger.error(f"Error parsing PDF file {filename}: {e}")
        raise ValueError(f"Could not parse PDF file: {e}")


def parse_docx_file(file_content: bytes, filename: str) -> str:
    """Extract text from DOCX file"""
    if not DOCX_AVAILABLE:
        raise ValueError("DOCX parsing not available. Install python-docx.")
    
    try:
        buffer = io.BytesIO(file_content)
        doc = Document(buffer)
        
        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # Also get text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    text_parts.append(" | ".join(row_text))
        
        full_text = "\n".join(text_parts)
        logger.info(f"Extracted {len(full_text)} characters from DOCX {filename}")
        return full_text
    except Exception as e:
        logger.error(f"Error parsing DOCX file {filename}: {e}")
        raise ValueError(f"Could not parse DOCX file: {e}")


# =============================================================================
# DATA EXTRACTION FROM EXCEL
# =============================================================================

def extract_company_info(sheets: Dict[str, pd.DataFrame]) -> Dict[str, Any]:
    """Extract company information from Summary_Statistics sheet"""
    info = {
        "company_name": "",
        "industry": "",
        "tax_year": 2024,
    }
    
    if "Summary_Statistics" in sheets:
        df = sheets["Summary_Statistics"]
        
        # Try to find company info by looking for Field/Value pattern
        columns = list(df.columns)
        
        for idx, row in df.iterrows():
            row_values = list(row.values)
            if len(row_values) >= 2:
                field = str(row_values[0]).lower().strip() if pd.notna(row_values[0]) else ""
                value = row_values[1] if len(row_values) > 1 and pd.notna(row_values[1]) else ""
                
                # Match specific fields
                if "company name" in field or field == "company":
                    if value and str(value).strip():
                        info["company_name"] = str(value).strip()
                elif "industry" in field:
                    if value and str(value).strip():
                        info["industry"] = str(value).strip()
                elif "tax year" in field or "year" in field:
                    try:
                        info["tax_year"] = int(float(value))
                    except (ValueError, TypeError):
                        pass
        
        # Fallback: look for company name patterns in values
        if not info["company_name"]:
            for idx, row in df.head(15).iterrows():
                for val in row.values:
                    if pd.notna(val):
                        val_str = str(val).strip()
                        # Look for typical company name patterns
                        if any(pattern in val_str.lower() for pattern in [" inc", " llc", " corp", " ltd", " solutions", " tech"]):
                            if len(val_str) > 5 and "company name" not in val_str.lower():
                                info["company_name"] = val_str
                                break
                if info["company_name"]:
                    break
    
    return info


def extract_projects(sheets: Dict[str, pd.DataFrame]) -> List[RDProject]:
    """Extract projects from Projects sheet"""
    projects = []
    
    if "Projects" not in sheets:
        return projects
    
    df = sheets["Projects"]
    
    for idx, row in df.iterrows():
        try:
            project = RDProject(
                project_id=str(row.get("Project_ID", f"P{idx}")),
                project_name=str(row.get("Project_Name", f"Project {idx}")),
                category=str(row.get("Category", "")) if pd.notna(row.get("Category")) else None,
                description=str(row.get("Description", "")) if pd.notna(row.get("Description")) else None,
                budget=float(row.get("Budget", 0)) if pd.notna(row.get("Budget")) else None,
            )
            
            # Extract existing four-part test flags if present
            fpt = FourPartTestResult()
            
            if pd.notna(row.get("Permitted_Purpose")):
                val = str(row.get("Permitted_Purpose")).lower()
                if val in ["true", "yes", "1", "pass"]:
                    fpt.permitted_purpose = TestStatus.PASS
                elif val in ["false", "no", "0", "fail"]:
                    fpt.permitted_purpose = TestStatus.FAIL
                    
            if pd.notna(row.get("Elimination_Uncertainty")):
                val = str(row.get("Elimination_Uncertainty")).lower()
                if val in ["true", "yes", "1", "pass"]:
                    fpt.elimination_uncertainty = TestStatus.PASS
                elif val in ["false", "no", "0", "fail"]:
                    fpt.elimination_uncertainty = TestStatus.FAIL
                    
            if pd.notna(row.get("Process_Experimentation")):
                val = str(row.get("Process_Experimentation")).lower()
                if val in ["true", "yes", "1", "pass"]:
                    fpt.process_experimentation = TestStatus.PASS
                elif val in ["false", "no", "0", "fail"]:
                    fpt.process_experimentation = TestStatus.FAIL
                    
            if pd.notna(row.get("Technological_Nature")):
                val = str(row.get("Technological_Nature")).lower()
                if val in ["true", "yes", "1", "pass"]:
                    fpt.technological_nature = TestStatus.PASS
                elif val in ["false", "no", "0", "fail"]:
                    fpt.technological_nature = TestStatus.FAIL
            
            project.four_part_test = fpt
            project.qualified = fpt.pass_count == 4
            
            projects.append(project)
        except Exception as e:
            logger.warning(f"Error extracting project at row {idx}: {e}")
    
    logger.info(f"Extracted {len(projects)} projects")
    return projects


def extract_employees(sheets: Dict[str, pd.DataFrame]) -> List[RDEmployee]:
    """Extract employees from Employees sheet"""
    employees = []
    
    if "Employees" not in sheets:
        return employees
    
    df = sheets["Employees"]
    
    for idx, row in df.iterrows():
        try:
            employee = RDEmployee(
                employee_id=str(row.get("Employee_ID", f"E{idx}")),
                name=str(row.get("Name", f"Employee {idx}")),
                job_title=str(row.get("Job_Title", "")) if pd.notna(row.get("Job_Title")) else None,
                department=str(row.get("Department", "")) if pd.notna(row.get("Department")) else None,
                location=str(row.get("Location", row.get("Location_State", ""))) if pd.notna(row.get("Location", row.get("Location_State"))) else None,
                w2_wages=float(row.get("W2_Wages", row.get("Box1_Wages", 0))) if pd.notna(row.get("W2_Wages", row.get("Box1_Wages"))) else 0.0,
                qre_wage_base=float(row.get("QRE_Wage_Base", 0)) if pd.notna(row.get("QRE_Wage_Base")) else 0.0,
                rd_allocation_percent=float(row.get("RD_Allocation_%", row.get("RD_Allocation_Percent", 0))) if pd.notna(row.get("RD_Allocation_%", row.get("RD_Allocation_Percent"))) else 0.0,
                stock_compensation=float(row.get("Stock_Compensation", 0)) if pd.notna(row.get("Stock_Compensation")) else 0.0,
                severance=float(row.get("Severance", 0)) if pd.notna(row.get("Severance")) else 0.0,
            )
            employees.append(employee)
        except Exception as e:
            logger.warning(f"Error extracting employee at row {idx}: {e}")
    
    logger.info(f"Extracted {len(employees)} employees")
    return employees


def extract_vendors(sheets: Dict[str, pd.DataFrame]) -> List[RDVendor]:
    """Extract vendors from Vendors sheet"""
    vendors = []
    
    if "Vendors" not in sheets:
        return vendors
    
    df = sheets["Vendors"]
    
    for idx, row in df.iterrows():
        try:
            risk_bearer = str(row.get("Risk_Bearer", "")).strip()
            ip_rights = str(row.get("IP_Rights", "")).strip()
            
            # Check if vendor qualifies for Sec 41 (company bears risk, company/shared IP)
            qualified = (
                risk_bearer.lower() in ["company", "taxpayer"] and
                ip_rights.lower() in ["company", "shared"]
            )
            
            vendor = RDVendor(
                vendor_id=str(row.get("Vendor_ID", f"V{idx}")),
                vendor_name=str(row.get("Vendor_Name", f"Vendor {idx}")),
                risk_bearer=risk_bearer,
                ip_rights=ip_rights,
                country=str(row.get("Country", row.get("Location", ""))) if pd.notna(row.get("Country", row.get("Location"))) else "",
                qualified=qualified,
            )
            vendors.append(vendor)
        except Exception as e:
            logger.warning(f"Error extracting vendor at row {idx}: {e}")
    
    logger.info(f"Extracted {len(vendors)} vendors")
    return vendors


def extract_expenses(sheets: Dict[str, pd.DataFrame], vendors: List[RDVendor]) -> List[RDExpense]:
    """Extract expenses from AP_Transactions and Supplies sheets"""
    expenses = []
    vendor_qualified = {v.vendor_id: v.qualified for v in vendors}
    
    # Contract research from AP_Transactions
    if "AP_Transactions" in sheets:
        df = sheets["AP_Transactions"]
        for idx, row in df.iterrows():
            try:
                vendor_id = str(row.get("Vendor_ID", ""))
                qualified = row.get("Qualified_Contract_Research", False)
                if isinstance(qualified, str):
                    qualified = qualified.lower() in ["true", "yes", "1"]
                
                # Also check vendor qualification
                if vendor_id and vendor_id in vendor_qualified:
                    qualified = qualified and vendor_qualified[vendor_id]
                
                expense = RDExpense(
                    transaction_id=str(row.get("Transaction_ID", f"AP{idx}")),
                    vendor_id=vendor_id if vendor_id else None,
                    description=str(row.get("Description", "")),
                    amount=float(row.get("Amount", 0)) if pd.notna(row.get("Amount")) else 0.0,
                    qre_amount=float(row.get("QRE_Amount", 0)) if pd.notna(row.get("QRE_Amount")) else 0.0,
                    qualified=qualified,
                    category="contract_research",
                )
                expenses.append(expense)
            except Exception as e:
                logger.warning(f"Error extracting AP transaction at row {idx}: {e}")
    
    # Supplies
    if "Supplies" in sheets:
        df = sheets["Supplies"]
        for idx, row in df.iterrows():
            try:
                qualified = row.get("Qualified_Supply", False)
                if isinstance(qualified, str):
                    qualified = qualified.lower() in ["true", "yes", "1"]
                
                expense = RDExpense(
                    transaction_id=str(row.get("Supply_ID", f"S{idx}")),
                    description=str(row.get("Item_Description", row.get("Description", ""))),
                    amount=float(row.get("Total_Amount", 0)) if pd.notna(row.get("Total_Amount")) else 0.0,
                    qre_amount=float(row.get("QRE_Amount", 0)) if pd.notna(row.get("QRE_Amount")) else 0.0,
                    qualified=qualified,
                    category="supplies",
                )
                expenses.append(expense)
            except Exception as e:
                logger.warning(f"Error extracting supply at row {idx}: {e}")
    
    logger.info(f"Extracted {len(expenses)} expenses")
    return expenses


def extract_qre_summary(sheets: Dict[str, pd.DataFrame]) -> Dict[str, float]:
    """Extract QRE summary totals from dedicated sheet if present"""
    summary = {
        "wage_qre": 0.0,
        "supply_qre": 0.0,
        "contract_qre": 0.0,
        "total_qre": 0.0,
    }
    
    if "QRE_Summary_2024" in sheets:
        df = sheets["QRE_Summary_2024"]
        if len(df) > 0:
            row = df.iloc[0]
            summary["wage_qre"] = float(row.get("Wage_QRE_2024", 0)) if pd.notna(row.get("Wage_QRE_2024")) else 0.0
            summary["supply_qre"] = float(row.get("Supplies_QRE_2024", 0)) if pd.notna(row.get("Supplies_QRE_2024")) else 0.0
            summary["contract_qre"] = float(row.get("Contract_QRE_2024", 0)) if pd.notna(row.get("Contract_QRE_2024")) else 0.0
            summary["total_qre"] = float(row.get("Total_QRE_2024", 0)) if pd.notna(row.get("Total_QRE_2024")) else 0.0
    
    return summary


def calculate_qre_from_data(
    employees: List[RDEmployee],
    expenses: List[RDExpense],
    vendors: List[RDVendor]
) -> Dict[str, float]:
    """
    Calculate QRE (Qualified Research Expenses) from raw employee and expense data.
    
    QRE Categories:
    1. Wage QRE = Sum of (W2 Wages - Stock Comp - Severance) * R&D Allocation %
       OR if QRE_Wage_Base is provided, use that instead
    2. Supply QRE = Sum of qualified supply expenses
    3. Contract QRE = 65% of qualified contract research expenses (per IRC Sec 41)
    """
    summary = {
        "wage_qre": 0.0,
        "supply_qre": 0.0,
        "contract_qre": 0.0,
        "total_qre": 0.0,
    }
    
    # Calculate Wage QRE
    for emp in employees:
        if emp.rd_allocation_percent > 0:
            # Use QRE Wage Base if provided, otherwise calculate from W2
            if emp.qre_wage_base > 0:
                wage_base = emp.qre_wage_base
            else:
                # QRE wage base = W2 wages minus excluded compensation
                wage_base = emp.w2_wages - emp.stock_compensation - emp.severance
                wage_base = max(0, wage_base)  # Ensure non-negative
            
            # Apply R&D allocation percentage
            emp_qre = wage_base * (emp.rd_allocation_percent / 100)
            summary["wage_qre"] += emp_qre
    
    # Calculate Supply QRE
    for exp in expenses:
        if exp.category == "supplies":
            if exp.qre_amount > 0:
                # Use pre-calculated QRE amount if provided
                summary["supply_qre"] += exp.qre_amount
            elif exp.qualified:
                # Otherwise use full amount if qualified
                summary["supply_qre"] += exp.amount
    
    # Calculate Contract Research QRE (65% per IRC Sec 41)
    CONTRACT_QRE_RATE = 0.65
    for exp in expenses:
        if exp.category == "contract_research":
            if exp.qre_amount > 0:
                # Use pre-calculated QRE amount if provided
                summary["contract_qre"] += exp.qre_amount
            elif exp.qualified:
                # Otherwise apply 65% rate to qualified amount
                summary["contract_qre"] += exp.amount * CONTRACT_QRE_RATE
    
    # Calculate total
    summary["total_qre"] = summary["wage_qre"] + summary["supply_qre"] + summary["contract_qre"]
    
    logger.info(f"Calculated QRE: Wage=${summary['wage_qre']:,.2f}, Supply=${summary['supply_qre']:,.2f}, Contract=${summary['contract_qre']:,.2f}, Total=${summary['total_qre']:,.2f}")
    
    return summary


# =============================================================================
# GAP ANALYSIS
# =============================================================================

def identify_gaps(session: RDAnalysisSession) -> List[GapItem]:
    """Identify information gaps that need user attention"""
    gaps = []
    
    # Check projects for missing four-part test data
    for project in session.projects:
        fpt = project.four_part_test
        missing = []
        
        if fpt.permitted_purpose == TestStatus.MISSING_DATA:
            missing.append("Permitted Purpose documentation")
        if fpt.elimination_uncertainty == TestStatus.MISSING_DATA:
            missing.append("Elimination of Uncertainty evidence")
        if fpt.process_experimentation == TestStatus.MISSING_DATA:
            missing.append("Process of Experimentation documentation")
        if fpt.technological_nature == TestStatus.MISSING_DATA:
            missing.append("Technological Nature justification")
        
        if missing:
            gaps.append(GapItem(
                gap_id=f"gap-{project.project_id}",
                category="project",
                item_id=project.project_id,
                item_name=project.project_name,
                gap_type="missing_data",
                description=f"Project needs additional documentation for four-part test",
                required_info=missing,
                priority="high" if len(missing) >= 3 else "medium",
            ))
        
        # Check for needs_review items
        needs_review = []
        if fpt.permitted_purpose == TestStatus.NEEDS_REVIEW:
            needs_review.append("Permitted Purpose needs clarification")
        if fpt.elimination_uncertainty == TestStatus.NEEDS_REVIEW:
            needs_review.append("Uncertainty elimination needs clarification")
        if fpt.process_experimentation == TestStatus.NEEDS_REVIEW:
            needs_review.append("Experimentation process needs clarification")
        if fpt.technological_nature == TestStatus.NEEDS_REVIEW:
            needs_review.append("Technological basis needs clarification")
        
        if needs_review:
            gaps.append(GapItem(
                gap_id=f"gap-review-{project.project_id}",
                category="project",
                item_id=project.project_id,
                item_name=project.project_name,
                gap_type="needs_clarification",
                description=f"Project needs additional clarification",
                required_info=needs_review,
                priority="medium",
            ))
    
    # Check vendors for risk/IP issues
    for vendor in session.vendors:
        if not vendor.qualified:
            gaps.append(GapItem(
                gap_id=f"gap-{vendor.vendor_id}",
                category="vendor",
                item_id=vendor.vendor_id,
                item_name=vendor.vendor_name,
                gap_type="needs_clarification",
                description=f"Vendor does not meet Sec.41 risk/IP requirements (Risk: {vendor.risk_bearer}, IP: {vendor.ip_rights})",
                required_info=["Contract showing company bears economic risk", "Documentation of IP ownership"],
                priority="medium",
            ))
    
    # Check employees with high stock comp or severance
    for emp in session.employees:
        issues = []
        if emp.w2_wages > 0:
            if emp.stock_compensation > emp.w2_wages * 0.5 or emp.stock_compensation > 200000:
                issues.append(f"Stock compensation ({emp.stock_compensation:,.0f}) may need review")
            if emp.severance > emp.w2_wages * 0.4:
                issues.append(f"Severance ({emp.severance:,.0f}) exceeds 40% of wages")
        
        if issues:
            gaps.append(GapItem(
                gap_id=f"gap-emp-{emp.employee_id}",
                category="employee",
                item_id=emp.employee_id,
                item_name=emp.name,
                gap_type="needs_clarification",
                description="Compensation items need review",
                required_info=issues,
                priority="low",
            ))
    
    logger.info(f"Identified {len(gaps)} gaps")
    return gaps


# =============================================================================
# GEMINI AI INTEGRATION
# =============================================================================

# Strict IRS Section 41 Four-Part Test Definitions
FOUR_PART_TEST_DEFINITIONS = """
## IRS SECTION 41 FOUR-PART TEST - STRICT EVALUATION CRITERIA

You are an expert R&D tax credit analyst. Evaluate STRICTLY according to IRS guidelines.
Be conservative - when in doubt, mark as "needs_review" not "pass".

### TEST 1: PERMITTED PURPOSE (IRC Section 174)

**PASS** - Activity MUST aim to:
- Develop NEW functionality that did not previously exist
- IMPROVE existing function, performance, reliability, or quality (not just cosmetic)
- Create a new or improved business component (product, process, technique, formula, invention, software)
- The improvement must be MORE than incremental/routine

**FAIL** - Activity is:
- Routine data collection or quality control
- Cosmetic or aesthetic changes only
- Adapting existing product to customer specifications WITHOUT technical uncertainty
- Style, taste, or seasonal design changes
- Market research or surveys
- Management studies or efficiency surveys
- Advertising or promotions
- Acquisition of another's patent or know-how

**NEEDS_REVIEW** - Unclear if improvement is substantial or routine

### TEST 2: ELIMINATION OF UNCERTAINTY

**PASS** - At project START, there was genuine uncertainty about:
- CAPABILITY: Can it be done at all? Is it technically feasible?
- METHOD: What is the best approach? How should it be developed?
- DESIGN: What is the appropriate design? What specifications will work?

The uncertainty must be TECHNOLOGICAL, not just business/economic uncertainty.

**FAIL** - Activity where:
- Outcome was known or predictable from the start
- Following well-established procedures or industry standards
- Only uncertainty was about cost, time, or market acceptance
- Using proven technology in standard applications
- Implementing vendor's documented solution

**NEEDS_REVIEW** - Some technical questions but unclear if truly uncertain

### TEST 3: PROCESS OF EXPERIMENTATION

**PASS** - Taxpayer MUST have used one or more of:
- Systematic trial and error (documented attempts and failures)
- Modeling or simulation to test alternatives
- Testing of hypotheses with measurable outcomes
- Iterative design with documented iterations
- Evaluation of alternatives through technical analysis

Evidence should show: hypothesis → test → evaluate → refine cycle

**FAIL** - Activity where:
- No documented testing or evaluation occurred
- Single approach used with no alternatives considered
- Purchased off-the-shelf solution
- Followed step-by-step instructions without modification
- "Trial and error" was just guessing without systematic approach

**NEEDS_REVIEW** - Some testing occurred but documentation is unclear

### TEST 4: TECHNOLOGICAL IN NATURE

**PASS** - Process fundamentally relies on principles of:
- Physical sciences (physics, chemistry, materials science)
- Biological sciences (biology, biochemistry, microbiology)
- Engineering disciplines (mechanical, electrical, civil, chemical, software)
- Computer science (algorithms, data structures, system architecture)

**FAIL** - Process relies primarily on:
- Social sciences, economics, psychology
- Business or management studies
- Market research or consumer behavior
- Artistic or aesthetic development
- Human factors without engineering basis

**NEEDS_REVIEW** - Mixed technical/non-technical basis
"""


# Model configuration - use gemini-3-flash-preview (latest as of Jan 2026)
RD_MODEL_NAME = os.environ.get("GEMINI_MODEL", "gemini-3-flash-preview")

# Initialize model
_rd_model = None
_gemini_configured = False

def reset_gemini_model():
    """Reset the Gemini model singleton to allow re-initialization"""
    global _rd_model, _gemini_configured
    _rd_model = None
    _gemini_configured = False
    logger.info("Gemini model singleton reset")

def _get_gemini_model():
    """Get or create Gemini model (singleton pattern, retries on error)"""
    global _rd_model, _gemini_configured
    
    # #region agent log
    _dbg_path = "/Users/dhruvramasubban/Desktop/TaxScapeCursor/.cursor/debug.log"
    def _dbg(loc, msg, data):
        try:
            import json as _j
            with open(_dbg_path, "a") as f:
                f.write(_j.dumps({"location": loc, "message": msg, "data": data, "timestamp": int(__import__('time').time()*1000), "sessionId": "debug-session", "hypothesisId": "A,B,C"}) + "\n")
        except: pass
        logger.info(f"[DEBUG] {loc}: {msg} | {data}")
    # #endregion
    
    # #region agent log
    _dbg("rd_parser:_get_gemini_model:entry", "Starting model init", {"cached": _rd_model is not None, "configured": _gemini_configured, "model_name": RD_MODEL_NAME})
    # #endregion
    
    # Return cached model if available
    if _rd_model is not None:
        return _rd_model
    
    # Check dependencies
    if not GEMINI_AVAILABLE:
        # #region agent log
        _dbg("rd_parser:_get_gemini_model:no_sdk", "SDK not available", {"GEMINI_AVAILABLE": False})
        # #endregion
        raise ValueError("google-generativeai package not installed. Run: pip install google-generativeai")
    
    # Check API key
    api_key = os.environ.get("GOOGLE_CLOUD_API_KEY") or os.environ.get("GEMINI_API_KEY") or os.environ.get("GOOGLE_API_KEY")
    
    # #region agent log
    _dbg("rd_parser:_get_gemini_model:api_key_check", "Checking API key", {"has_key": bool(api_key), "key_length": len(api_key) if api_key else 0, "key_prefix": api_key[:8] if api_key else None})
    # #endregion
    
    if not api_key:
        raise ValueError("GOOGLE_CLOUD_API_KEY or GEMINI_API_KEY not set")
    
    # Try to initialize model (don't cache errors - allow retry)
    try:
        # #region agent log
        _dbg("rd_parser:_get_gemini_model:before_configure", "About to configure genai", {"model": RD_MODEL_NAME})
        # #endregion
        
        # Always reconfigure to ensure fresh state
        genai.configure(api_key=api_key)
        _gemini_configured = True
        
        # #region agent log
        _dbg("rd_parser:_get_gemini_model:after_configure", "genai configured, creating model", {"model": RD_MODEL_NAME})
        # #endregion
        
        _rd_model = genai.GenerativeModel(RD_MODEL_NAME)
        
        # #region agent log
        _dbg("rd_parser:_get_gemini_model:success", "Model created successfully", {"model": RD_MODEL_NAME})
        # #endregion
        
        logger.info(f"R&D Gemini model initialized: {RD_MODEL_NAME}")
        return _rd_model
    except Exception as e:
        # Reset state to allow retry
        _rd_model = None
        # #region agent log
        _dbg("rd_parser:_get_gemini_model:error", "Model init failed", {"error": str(e), "error_type": type(e).__name__})
        # #endregion
        logger.error(f"Failed to initialize Gemini model: {e}")
        raise ValueError(f"Failed to initialize Gemini model: {str(e)}")

# Backwards-compatible alias
def _get_gemini_client():
    """Backwards-compatible alias for _get_gemini_model"""
    return _get_gemini_model()


def generate_ai_content(prompt: str, temperature: float = 0.2, max_output_tokens: int = 4096) -> str:
    """Generate content using Gemini model (legacy SDK)"""
    # #region agent log
    _dbg_path = "/Users/dhruvramasubban/Desktop/TaxScapeCursor/.cursor/debug.log"
    import time as _time
    _start = _time.time()
    def _dbg(loc, msg, data):
        try:
            import json as _j
            with open(_dbg_path, "a") as f:
                f.write(_j.dumps({"location": loc, "message": msg, "data": data, "timestamp": int(_time.time()*1000), "sessionId": "debug-session", "hypothesisId": "B,C,D,E"}) + "\n")
        except: pass
        logger.info(f"[DEBUG] {loc}: {msg} | {data}")
    _dbg("rd_parser:generate_ai_content:entry", "Starting content generation", {"prompt_len": len(prompt), "temp": temperature, "max_tokens": max_output_tokens})
    # #endregion
    
    model = _get_gemini_model()
    
    # #region agent log
    _dbg("rd_parser:generate_ai_content:before_generate", "About to call generate_content", {"model": RD_MODEL_NAME})
    # #endregion
    
    try:
        response = model.generate_content(
            prompt,
            generation_config=genai.types.GenerationConfig(
                temperature=temperature,
                max_output_tokens=max_output_tokens,
            )
        )
    except Exception as gen_err:
        # #region agent log
        _dbg("rd_parser:generate_ai_content:generate_error", "generate_content threw exception", {"error": str(gen_err), "error_type": type(gen_err).__name__, "elapsed_ms": int((_time.time()-_start)*1000)})
        # #endregion
        raise
    
    # #region agent log
    _dbg("rd_parser:generate_ai_content:after_generate", "generate_content returned", {
        "has_response": response is not None,
        "has_candidates": hasattr(response, 'candidates') and bool(response.candidates) if response else False,
        "candidate_count": len(response.candidates) if response and hasattr(response, 'candidates') and response.candidates else 0,
        "finish_reason": response.candidates[0].finish_reason if response and hasattr(response, 'candidates') and response.candidates else None,
        "elapsed_ms": int((_time.time()-_start)*1000)
    })
    # #endregion
    
    # Handle blocked responses (safety filters)
    if response and hasattr(response, 'candidates') and response.candidates:
        candidate = response.candidates[0]
        if hasattr(candidate, 'finish_reason') and candidate.finish_reason == 2:
            # #region agent log
            _dbg("rd_parser:generate_ai_content:safety_blocked", "Response blocked by safety filters", {"finish_reason": 2})
            # #endregion
            raise ValueError("AI response blocked by safety filters. Please rephrase your request.")
    
    # Try to get text, handle potential errors
    try:
        if response and hasattr(response, 'text') and response.text:
            # #region agent log
            _dbg("rd_parser:generate_ai_content:success", "Got response text", {"text_len": len(response.text), "elapsed_ms": int((_time.time()-_start)*1000)})
            # #endregion
            return response.text
    except ValueError as e:
        # #region agent log
        _dbg("rd_parser:generate_ai_content:text_access_error", "Error accessing response.text", {"error": str(e), "elapsed_ms": int((_time.time()-_start)*1000)})
        # #endregion
        # This happens when response.text is accessed but no valid parts exist
        if "finish_reason" in str(e) or "Part" in str(e):
            raise ValueError("AI response was empty or blocked. Please try again.")
        raise
    
    # #region agent log
    _dbg("rd_parser:generate_ai_content:empty_response", "Response was empty", {"elapsed_ms": int((_time.time()-_start)*1000)})
    # #endregion
    raise ValueError("AI returned empty response")


def check_ai_available() -> Dict[str, Any]:
    """Check if AI is available and configured"""
    result = {
        "available": False,
        "gemini_installed": GEMINI_AVAILABLE,
        "api_key_set": False,
        "error": None,
        "model": RD_MODEL_NAME
    }
    
    if not GEMINI_AVAILABLE:
        result["error"] = "google-genai package not installed"
        return result
    
    api_key = os.environ.get("GOOGLE_CLOUD_API_KEY") or os.environ.get("GEMINI_API_KEY") or os.environ.get("GOOGLE_API_KEY")
    if not api_key:
        result["error"] = "GOOGLE_CLOUD_API_KEY or GEMINI_API_KEY environment variable not set"
        return result
    
    result["api_key_set"] = True
    
    try:
        # Just verify the model can be initialized - skip content generation test
        # Content generation tests can trigger unpredictable safety filters
        model = _get_gemini_model()
        if model:
            result["available"] = True
            logger.info("R&D AI check passed")
        else:
            result["error"] = "Failed to initialize AI model"
    except Exception as e:
        result["error"] = f"AI connection test failed: {str(e)}"
        logger.error(f"R&D AI check failed: {e}")
    
    return result


def parse_gap_documents(files: List[Dict[str, Any]]) -> str:
    """Parse uploaded gap documents and extract text content"""
    extracted_texts = []
    
    for file_info in files:
        filename = file_info.get("filename", "unknown")
        content = file_info.get("content", b"")
        
        try:
            if filename.lower().endswith(".pdf"):
                if PDF_AVAILABLE:
                    text = parse_pdf_file(content, filename)
                    extracted_texts.append(f"=== From {filename} ===\n{text}")
                else:
                    extracted_texts.append(f"=== {filename} ===\n[PDF parsing not available]")
                    
            elif filename.lower().endswith(".docx"):
                if DOCX_AVAILABLE:
                    text = parse_docx_file(content, filename)
                    extracted_texts.append(f"=== From {filename} ===\n{text}")
                else:
                    extracted_texts.append(f"=== {filename} ===\n[DOCX parsing not available]")
                    
            elif filename.lower().endswith((".xlsx", ".xls")):
                # Extract text from Excel
                sheets = parse_excel_file(content, filename)
                text_parts = []
                for sheet_name, df in sheets.items():
                    text_parts.append(f"Sheet: {sheet_name}")
                    text_parts.append(df.to_string(max_rows=50))
                extracted_texts.append(f"=== From {filename} ===\n" + "\n".join(text_parts))
                
            elif filename.lower().endswith((".txt", ".csv")):
                text = content.decode("utf-8", errors="ignore")
                extracted_texts.append(f"=== From {filename} ===\n{text[:5000]}")
                
        except Exception as e:
            logger.error(f"Error parsing gap document {filename}: {e}")
            extracted_texts.append(f"=== {filename} ===\n[Error parsing: {str(e)}]")
    
    return "\n\n".join(extracted_texts)


def evaluate_project_with_ai(project: RDProject, additional_context: str = "") -> RDProject:
    """Use Gemini to evaluate a project against the four-part test with strict IRS criteria"""
    try:
        client = _get_gemini_client()
        
        prompt = f"""{FOUR_PART_TEST_DEFINITIONS}

---

## PROJECT TO EVALUATE

**Project Name:** {project.project_name}
**Category:** {project.category or 'Not specified'}
**Description:** {project.description or 'No description provided - this is a significant gap'}
**Budget:** {f'${project.budget:,.2f}' if project.budget else 'Not specified'}

{('## ADDITIONAL DOCUMENTATION/CONTEXT' + chr(10) + additional_context) if additional_context else '## NO ADDITIONAL DOCUMENTATION PROVIDED' + chr(10) + 'Without supporting documentation, many tests will likely need review.'}

---

## YOUR TASK

Evaluate this project against each of the four tests using the STRICT criteria above.

For each test:
1. Cite specific evidence from the project description/documentation
2. If evidence is missing, say what specific documentation would be needed
3. Be CONSERVATIVE - if evidence is weak or missing, mark as "needs_review" or "fail"

Respond in this EXACT JSON format (no markdown, just raw JSON):
{{
    "permitted_purpose": {{
        "status": "pass" | "fail" | "needs_review",
        "reasoning": "Cite specific evidence or explain what's missing. Be specific."
    }},
    "elimination_uncertainty": {{
        "status": "pass" | "fail" | "needs_review",
        "reasoning": "What technical uncertainty existed? Cite evidence or explain gap."
    }},
    "process_experimentation": {{
        "status": "pass" | "fail" | "needs_review",
        "reasoning": "What experimentation was performed? Cite evidence or explain gap."
    }},
    "technological_nature": {{
        "status": "pass" | "fail" | "needs_review",
        "reasoning": "What scientific/engineering principles were applied? Cite evidence."
    }},
    "confidence_score": 0.0 to 1.0,
    "summary": "2-3 sentence overall assessment of qualification likelihood",
    "missing_info": ["specific document or information needed to improve evaluation"]
}}
"""
        
        logger.info(f"Sending project {project.project_id} to AI for evaluation using {RD_MODEL_NAME}...")
        
        # Use helper function for legacy SDK
        response_text = generate_ai_content(prompt, temperature=0.2, max_output_tokens=4096)
        response_text = response_text.strip()
        logger.info(f"AI response received for project {project.project_id}, length: {len(response_text)}")
        
        # Extract JSON from response (handle markdown code blocks)
        if "```json" in response_text:
            response_text = response_text.split("```json")[1].split("```")[0]
        elif "```" in response_text:
            response_text = response_text.split("```")[1].split("```")[0]
        
        response_text = response_text.strip()
        
        try:
            result = json.loads(response_text)
        except json.JSONDecodeError as je:
            logger.error(f"JSON parse error: {je}. Response was: {response_text[:500]}")
            raise ValueError(f"AI response was not valid JSON: {str(je)}")
        
        # Update project with AI evaluation
        fpt = project.four_part_test
        
        status_map = {
            "pass": TestStatus.PASS,
            "fail": TestStatus.FAIL,
            "needs_review": TestStatus.NEEDS_REVIEW,
        }
        
        if "permitted_purpose" in result:
            status = result["permitted_purpose"].get("status", "needs_review").lower()
            fpt.permitted_purpose = status_map.get(status, TestStatus.NEEDS_REVIEW)
            fpt.permitted_purpose_reasoning = result["permitted_purpose"].get("reasoning", "No reasoning provided")
        
        if "elimination_uncertainty" in result:
            status = result["elimination_uncertainty"].get("status", "needs_review").lower()
            fpt.elimination_uncertainty = status_map.get(status, TestStatus.NEEDS_REVIEW)
            fpt.elimination_uncertainty_reasoning = result["elimination_uncertainty"].get("reasoning", "No reasoning provided")
        
        if "process_experimentation" in result:
            status = result["process_experimentation"].get("status", "needs_review").lower()
            fpt.process_experimentation = status_map.get(status, TestStatus.NEEDS_REVIEW)
            fpt.process_experimentation_reasoning = result["process_experimentation"].get("reasoning", "No reasoning provided")
        
        if "technological_nature" in result:
            status = result["technological_nature"].get("status", "needs_review").lower()
            fpt.technological_nature = status_map.get(status, TestStatus.NEEDS_REVIEW)
            fpt.technological_nature_reasoning = result["technological_nature"].get("reasoning", "No reasoning provided")
        
        project.four_part_test = fpt
        project.confidence_score = float(result.get("confidence_score", 0.5))
        project.ai_summary = result.get("summary", "Evaluation completed")
        project.missing_info = result.get("missing_info", [])
        project.qualified = fpt.pass_count == 4
        
        logger.info(f"AI evaluated project {project.project_id}: {fpt.pass_count}/4 tests pass, qualified={project.qualified}")
        
    except Exception as e:
        error_msg = str(e)
        logger.error(f"Error evaluating project {project.project_id} with AI: {error_msg}")
        
        # Set all tests to needs_review with error info
        project.four_part_test.permitted_purpose = TestStatus.NEEDS_REVIEW
        project.four_part_test.permitted_purpose_reasoning = f"AI evaluation error: {error_msg}"
        project.four_part_test.elimination_uncertainty = TestStatus.NEEDS_REVIEW
        project.four_part_test.elimination_uncertainty_reasoning = f"AI evaluation error: {error_msg}"
        project.four_part_test.process_experimentation = TestStatus.NEEDS_REVIEW
        project.four_part_test.process_experimentation_reasoning = f"AI evaluation error: {error_msg}"
        project.four_part_test.technological_nature = TestStatus.NEEDS_REVIEW
        project.four_part_test.technological_nature_reasoning = f"AI evaluation error: {error_msg}"
        
        project.ai_summary = f"AI evaluation failed: {error_msg}. Manual review required."
        project.confidence_score = 0.0
        project.missing_info = ["AI evaluation failed - please retry or review manually"]
        project.qualified = False
    
    return project


def re_evaluate_project_with_gap_context(
    project: RDProject, 
    gap_documents: List[Dict[str, Any]],
    existing_context: str = ""
) -> RDProject:
    """Re-evaluate a project after new gap documentation is uploaded"""
    
    # Parse the new gap documents
    new_context = parse_gap_documents(gap_documents)
    
    # Combine with existing context
    full_context = ""
    if existing_context:
        full_context += f"=== PREVIOUS CONTEXT ===\n{existing_context}\n\n"
    if new_context:
        full_context += f"=== NEW DOCUMENTATION UPLOADED ===\n{new_context}"
    
    logger.info(f"Re-evaluating project {project.project_id} with {len(gap_documents)} new documents")
    
    # Re-run AI evaluation with combined context
    return evaluate_project_with_ai(project, full_context)


def analyze_document_with_ai(text: str, context: str = "") -> Dict[str, Any]:
    """Use Gemini to extract structured information from document text"""
    try:
        client = _get_gemini_client()
        
        prompt = f"""Analyze this document and extract relevant R&D tax credit information.

DOCUMENT TEXT:
{text[:10000]}

{f"CONTEXT: {context}" if context else ""}

Extract any information related to:
1. R&D projects and their technical details
2. Experimentation processes used
3. Technical uncertainties being addressed
4. Scientific/engineering principles applied

Respond in JSON format:
{{
    "projects_mentioned": [
        {{
            "name": "project name",
            "description": "description",
            "technical_goal": "goal",
            "uncertainty": "uncertainty addressed",
            "experimentation": "methods used"
        }}
    ],
    "key_findings": ["list of relevant findings"],
    "four_part_test_evidence": {{
        "permitted_purpose": "evidence found",
        "elimination_uncertainty": "evidence found",
        "process_experimentation": "evidence found",
        "technological_nature": "evidence found"
    }}
}}
"""
        
        response_text = generate_ai_content(prompt, temperature=0.2, max_output_tokens=4096)
        
        # Extract JSON from response
        if "```json" in response_text:
            response_text = response_text.split("```json")[1].split("```")[0]
        elif "```" in response_text:
            response_text = response_text.split("```")[1].split("```")[0]
        
        return json.loads(response_text.strip())
        
    except Exception as e:
        logger.error(f"Error analyzing document with AI: {e}")
        return {"error": str(e)}


# =============================================================================
# MAIN ANALYSIS FUNCTION
# =============================================================================

def create_analysis_session(
    files: List[Dict[str, Any]],  # List of {"filename": str, "content": bytes, "content_type": str}
    use_ai: bool = True
) -> RDAnalysisSession:
    """Create and run a complete R&D analysis session"""
    
    session = RDAnalysisSession(
        session_id=str(uuid.uuid4()),
        created_at=datetime.utcnow().isoformat(),
    )
    
    all_sheets = {}
    document_texts = []
    
    # Parse all files
    for file_info in files:
        filename = file_info["filename"]
        content = file_info["content"]
        content_type = file_info.get("content_type", "")
        
        try:
            if filename.endswith((".xlsx", ".xls")):
                sheets = parse_excel_file(content, filename)
                all_sheets.update(sheets)
            elif filename.endswith(".pdf"):
                text = parse_pdf_file(content, filename)
                document_texts.append({"filename": filename, "text": text})
            elif filename.endswith(".docx"):
                text = parse_docx_file(content, filename)
                document_texts.append({"filename": filename, "text": text})
            elif filename.endswith(".csv"):
                df = pd.read_csv(io.BytesIO(content))
                all_sheets[filename] = df
        except Exception as e:
            session.errors.append(f"Error parsing {filename}: {str(e)}")
            logger.error(f"Error parsing {filename}: {e}")
    
    # Extract data from Excel sheets
    if all_sheets:
        company_info = extract_company_info(all_sheets)
        session.company_name = company_info.get("company_name", "")
        session.industry = company_info.get("industry", "")
        session.tax_year = company_info.get("tax_year", 2024)
        
        session.projects = extract_projects(all_sheets)
        session.employees = extract_employees(all_sheets)
        session.vendors = extract_vendors(all_sheets)
        session.expenses = extract_expenses(all_sheets, session.vendors)
        
        # Try to get QRE from summary sheet first
        qre_summary = extract_qre_summary(all_sheets)
        
        # If QRE summary is empty/zero, calculate from raw data
        if qre_summary["total_qre"] == 0:
            logger.info("No QRE summary found in input - calculating from employee/expense data...")
            qre_summary = calculate_qre_from_data(
                session.employees, 
                session.expenses, 
                session.vendors
            )
        
        session.wage_qre = qre_summary["wage_qre"]
        session.supply_qre = qre_summary["supply_qre"]
        session.contract_qre = qre_summary["contract_qre"]
        session.total_qre = qre_summary["total_qre"]
        
        # Calculate summary stats
        session.total_employees = len(session.employees)
        session.rd_employees = len([e for e in session.employees if e.rd_allocation_percent > 0])
    
    session.parsing_complete = True
    
    # Run AI evaluation on projects if enabled
    if use_ai and GEMINI_AVAILABLE:
        # Combine document texts for context
        additional_context = ""
        if document_texts:
            additional_context = "\n\n".join([f"From {d['filename']}:\n{d['text'][:2000]}" for d in document_texts])
        
        for i, project in enumerate(session.projects):
            try:
                session.projects[i] = evaluate_project_with_ai(project, additional_context)
            except Exception as e:
                logger.error(f"AI evaluation failed for project {project.project_id}: {e}")
        
        session.qualified_projects = len([p for p in session.projects if p.qualified])
    
    # Identify gaps
    session.gaps = identify_gaps(session)
    
    session.analysis_complete = True
    
    return session




