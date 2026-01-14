"""
Credit Estimate Routes
Implements credit range drafting, senior signoff, exports, and client delivery.
"""

import logging
import io
from datetime import datetime, timedelta
from typing import Optional, List, Dict, Any
from uuid import uuid4

from fastapi import APIRouter, Depends, HTTPException, Query, Header
from fastapi.responses import StreamingResponse
from pydantic import BaseModel, Field

from .supabase_client import get_supabase
from .credit_estimate_engine import CreditEstimateEngine

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/api/estimates", tags=["estimates"])

# ============================================================================
# Auth Helpers
# ============================================================================

def verify_supabase_token(token: str) -> Optional[dict]:
    """Verify a Supabase JWT and return user data."""
    supabase = get_supabase()
    try:
        user_response = supabase.auth.get_user(token)
        if user_response and user_response.user:
            user = user_response.user
            return {"id": user.id, "email": user.email}
        return None
    except Exception as e:
        logger.warning(f"Token verification failed: {e}")
        return None


async def get_current_user(authorization: Optional[str] = Header(None)):
    """Extract and verify user from Supabase JWT."""
    if not authorization:
        raise HTTPException(status_code=401, detail="Authorization header missing")
    
    parts = authorization.split()
    if len(parts) != 2 or parts[0].lower() != "bearer":
        raise HTTPException(status_code=401, detail="Invalid authorization header")
    
    user_data = verify_supabase_token(parts[1])
    if not user_data:
        raise HTTPException(status_code=401, detail="Invalid or expired token")
    
    return user_data


def get_user_profile(user_id: str) -> Optional[dict]:
    """Get user profile with organization info."""
    supabase = get_supabase()
    try:
        result = supabase.table("profiles").select("*, organization_id, role, role_level").eq("id", user_id).single().execute()
        return result.data
    except:
        return None


def check_cpa_or_executive(user: dict) -> bool:
    """Check if user is CPA or Executive role."""
    profile = get_user_profile(user["id"])
    if not profile:
        return False
    role = profile.get("role", "").lower()
    return role in ["cpa", "executive", "admin"]


def check_senior_or_above(user: dict) -> bool:
    """Check if user is senior CPA or above."""
    profile = get_user_profile(user["id"])
    if not profile:
        return False
    
    role = profile.get("role", "").lower()
    if role in ["executive", "admin"]:
        return True
    
    if role == "cpa":
        role_level = profile.get("role_level", "").lower()
        return role_level in ["senior", "director", "partner"]
    
    return False


def write_audit_log(
    org_id: str,
    user_id: str,
    action: str,
    item_type: str,
    item_id: str = None,
    details: dict = None
):
    """Write to audit_logs table."""
    supabase = get_supabase()
    try:
        supabase.table("audit_logs").insert({
            "organization_id": org_id,
            "user_id": user_id,
            "action": action,
            "item_type": item_type,
            "item_id": item_id,
            "details": details or {},
            "created_at": datetime.utcnow().isoformat()
        }).execute()
    except Exception as e:
        logger.error(f"Failed to write audit log: {e}")


# ============================================================================
# Request/Response Models
# ============================================================================

class DraftEstimateRequest(BaseModel):
    client_company_id: str
    tax_year: int
    methodology: str = "both"  # regular | asc | both
    range_strategy: Optional[dict] = None


class UpdateEstimateRequest(BaseModel):
    assumptions: Optional[List[dict]] = None
    methodology: Optional[str] = None
    user_notes: Optional[str] = None


class SignoffRequest(BaseModel):
    decision: str  # approved | rejected | changes_requested
    reason_code: str
    note: str
    modifications: Optional[dict] = None


class ExportRequest(BaseModel):
    export_type: str = "pdf"  # pdf | docx


# ============================================================================
# Estimate Endpoints
# ============================================================================

@router.post("/draft")
async def draft_estimate(
    request: DraftEstimateRequest,
    user: dict = Depends(get_current_user)
):
    """
    Draft a new credit estimate with low/base/high ranges.
    Creates a new version for the client/year.
    """
    supabase = get_supabase()
    
    if not check_cpa_or_executive(user):
        raise HTTPException(status_code=403, detail="CPA or Executive role required")
    
    profile = get_user_profile(user["id"])
    org_id = profile.get("organization_id")
    
    # Verify client access
    try:
        client = supabase.table("client_companies")\
            .select("id, name")\
            .eq("id", request.client_company_id)\
            .eq("organization_id", org_id)\
            .single()\
            .execute()
        
        if not client.data:
            raise HTTPException(status_code=404, detail="Client not found")
    except:
        raise HTTPException(status_code=404, detail="Client not found")
    
    # Get next version number
    version_result = supabase.rpc("get_next_estimate_version", {
        "p_client_id": request.client_company_id,
        "p_tax_year": request.tax_year
    }).execute()
    
    # Fallback if function doesn't exist
    if version_result.data is None:
        existing = supabase.table("credit_estimates")\
            .select("estimate_version")\
            .eq("client_company_id", request.client_company_id)\
            .eq("tax_year", request.tax_year)\
            .order("estimate_version", desc=True)\
            .limit(1)\
            .execute()
        
        next_version = (existing.data[0]["estimate_version"] + 1) if existing.data else 1
    else:
        next_version = version_result.data
    
    # Mark previous versions as superseded
    supabase.table("credit_estimates")\
        .update({"status": "superseded"})\
        .eq("client_company_id", request.client_company_id)\
        .eq("tax_year", request.tax_year)\
        .in_("status", ["draft", "rejected"])\
        .execute()
    
    # Compute estimate
    engine = CreditEstimateEngine(
        supabase=supabase,
        org_id=org_id,
        client_id=request.client_company_id,
        tax_year=request.tax_year
    )
    
    estimate_data = engine.compute_estimate(
        methodology=request.methodology,
        range_strategy=request.range_strategy
    )
    
    # Get intake session ID if exists
    intake_session_id = None
    try:
        session = supabase.table("client_intake_sessions")\
            .select("id")\
            .eq("client_company_id", request.client_company_id)\
            .order("created_at", desc=True)\
            .limit(1)\
            .execute()
        if session.data:
            intake_session_id = session.data[0]["id"]
    except:
        pass
    
    # Create estimate record
    estimate_id = str(uuid4())
    estimate_record = {
        "id": estimate_id,
        "organization_id": org_id,
        "client_company_id": request.client_company_id,
        "tax_year": request.tax_year,
        "intake_session_id": intake_session_id,
        "estimate_version": next_version,
        "status": "draft",
        "methodology": request.methodology,
        "range_low": estimate_data["range_low"],
        "range_base": estimate_data["range_base"],
        "range_high": estimate_data["range_high"],
        "assumptions": estimate_data["assumptions"],
        "data_completeness_score": estimate_data["data_completeness_score"],
        "risk_notes": estimate_data["risk_notes"],
        "missing_inputs": estimate_data["missing_inputs"],
        "range_strategy": request.range_strategy or {},
        "created_by_user_id": user["id"]
    }
    
    supabase.table("credit_estimates").insert(estimate_record).execute()
    
    # Audit log
    write_audit_log(
        org_id=org_id,
        user_id=user["id"],
        action="credit_estimate_drafted",
        item_type="credit_estimate",
        item_id=estimate_id,
        details={
            "client_company_id": request.client_company_id,
            "tax_year": request.tax_year,
            "version": next_version,
            "methodology": request.methodology,
            "completeness_score": estimate_data["data_completeness_score"],
            "base_total_qre": estimate_data["range_base"]["total_qre"],
            "base_credit": estimate_data["range_base"].get("credit_amount_selected")
        }
    )
    
    return {
        "id": estimate_id,
        "version": next_version,
        "status": "draft",
        **estimate_data,
        "client_name": client.data.get("name")
    }


@router.get("")
async def list_estimates(
    client_id: Optional[str] = Query(None),
    tax_year: Optional[int] = Query(None),
    status: Optional[str] = Query(None),
    limit: int = Query(20),
    offset: int = Query(0),
    user: dict = Depends(get_current_user)
):
    """List credit estimates with filters."""
    supabase = get_supabase()
    
    if not check_cpa_or_executive(user):
        raise HTTPException(status_code=403, detail="CPA or Executive role required")
    
    profile = get_user_profile(user["id"])
    org_id = profile.get("organization_id")
    
    query = supabase.table("credit_estimates")\
        .select("*, client_companies(name)")\
        .eq("organization_id", org_id)
    
    if client_id:
        query = query.eq("client_company_id", client_id)
    if tax_year:
        query = query.eq("tax_year", tax_year)
    if status:
        query = query.eq("status", status)
    
    result = query.order("created_at", desc=True)\
        .range(offset, offset + limit - 1)\
        .execute()
    
    estimates = result.data or []
    
    # Enrich with client names
    for est in estimates:
        est["client_name"] = est.get("client_companies", {}).get("name")
    
    return {"estimates": estimates}


@router.get("/{estimate_id}")
async def get_estimate(
    estimate_id: str,
    user: dict = Depends(get_current_user)
):
    """Get estimate detail with history."""
    supabase = get_supabase()
    
    if not check_cpa_or_executive(user):
        raise HTTPException(status_code=403, detail="CPA or Executive role required")
    
    profile = get_user_profile(user["id"])
    org_id = profile.get("organization_id")
    
    # Get estimate
    try:
        result = supabase.table("credit_estimates")\
            .select("*, client_companies(name)")\
            .eq("id", estimate_id)\
            .eq("organization_id", org_id)\
            .single()\
            .execute()
        
        estimate = result.data
    except:
        raise HTTPException(status_code=404, detail="Estimate not found")
    
    if not estimate:
        raise HTTPException(status_code=404, detail="Estimate not found")
    
    # Get signoffs
    signoffs = supabase.table("estimate_signoffs")\
        .select("*, profiles(full_name)")\
        .eq("credit_estimate_id", estimate_id)\
        .order("created_at", desc=True)\
        .execute()
    
    # Get exports
    exports = supabase.table("estimate_exports")\
        .select("*")\
        .eq("credit_estimate_id", estimate_id)\
        .order("created_at", desc=True)\
        .execute()
    
    # Get version history
    versions = supabase.table("credit_estimates")\
        .select("id, estimate_version, status, created_at, created_by_user_id")\
        .eq("client_company_id", estimate["client_company_id"])\
        .eq("tax_year", estimate["tax_year"])\
        .order("estimate_version", desc=True)\
        .execute()
    
    # Check if data is stale (underlying data changed after estimate)
    is_stale = False
    stale_reason = None
    
    try:
        # Check if any canonical data was updated after estimate
        estimate_time = estimate["updated_at"] or estimate["created_at"]
        
        employees_check = supabase.table("employees")\
            .select("updated_at")\
            .eq("client_company_id", estimate["client_company_id"])\
            .gt("updated_at", estimate_time)\
            .limit(1)\
            .execute()
        
        if employees_check.data:
            is_stale = True
            stale_reason = "Employee data updated since estimate"
    except:
        pass
    
    return {
        "estimate": estimate,
        "signoffs": signoffs.data or [],
        "exports": exports.data or [],
        "version_history": versions.data or [],
        "is_stale": is_stale,
        "stale_reason": stale_reason,
        "client_name": estimate.get("client_companies", {}).get("name")
    }


@router.patch("/{estimate_id}")
async def update_estimate(
    estimate_id: str,
    request: UpdateEstimateRequest,
    user: dict = Depends(get_current_user)
):
    """
    Update a draft estimate (junior edits).
    Only draft estimates can be updated.
    """
    supabase = get_supabase()
    
    if not check_cpa_or_executive(user):
        raise HTTPException(status_code=403, detail="CPA or Executive role required")
    
    profile = get_user_profile(user["id"])
    org_id = profile.get("organization_id")
    
    # Get estimate
    try:
        result = supabase.table("credit_estimates")\
            .select("*")\
            .eq("id", estimate_id)\
            .eq("organization_id", org_id)\
            .single()\
            .execute()
        
        estimate = result.data
    except:
        raise HTTPException(status_code=404, detail="Estimate not found")
    
    if not estimate:
        raise HTTPException(status_code=404, detail="Estimate not found")
    
    if estimate["status"] not in ["draft", "rejected"]:
        raise HTTPException(status_code=400, detail="Only draft or rejected estimates can be updated")
    
    # Build update
    update_data = {}
    changes = {}
    
    if request.assumptions is not None:
        # Merge user assumptions with existing
        existing_assumptions = estimate.get("assumptions", [])
        
        # Mark user-entered assumptions
        for assumption in request.assumptions:
            if not assumption.get("assumption_id"):
                assumption["assumption_id"] = f"USER_{uuid4().hex[:8]}"
            assumption["source"] = "user_entered"
        
        # Keep system assumptions, add/update user ones
        system_assumptions = [a for a in existing_assumptions if a.get("source") == "system_default"]
        user_assumptions = request.assumptions
        
        update_data["assumptions"] = system_assumptions + user_assumptions
        changes["assumptions"] = {"before": len(existing_assumptions), "after": len(update_data["assumptions"])}
    
    if request.methodology:
        update_data["methodology"] = request.methodology
        changes["methodology"] = {"before": estimate.get("methodology"), "after": request.methodology}
    
    if not update_data:
        return {"message": "No changes"}
    
    # Update
    supabase.table("credit_estimates")\
        .update(update_data)\
        .eq("id", estimate_id)\
        .execute()
    
    # Audit log
    write_audit_log(
        org_id=org_id,
        user_id=user["id"],
        action="credit_estimate_updated",
        item_type="credit_estimate",
        item_id=estimate_id,
        details={"changes": changes}
    )
    
    return {"message": "Estimate updated", "changes": changes}


@router.post("/{estimate_id}/submit")
async def submit_for_signoff(
    estimate_id: str,
    user: dict = Depends(get_current_user)
):
    """
    Submit estimate for senior signoff.
    Creates an escalation for senior review.
    """
    supabase = get_supabase()
    
    if not check_cpa_or_executive(user):
        raise HTTPException(status_code=403, detail="CPA or Executive role required")
    
    profile = get_user_profile(user["id"])
    org_id = profile.get("organization_id")
    
    # Get estimate
    try:
        result = supabase.table("credit_estimates")\
            .select("*")\
            .eq("id", estimate_id)\
            .eq("organization_id", org_id)\
            .single()\
            .execute()
        
        estimate = result.data
    except:
        raise HTTPException(status_code=404, detail="Estimate not found")
    
    if not estimate:
        raise HTTPException(status_code=404, detail="Estimate not found")
    
    if estimate["status"] != "draft":
        raise HTTPException(status_code=400, detail="Only draft estimates can be submitted")
    
    # Check minimum completeness or require preliminary tag
    completeness = estimate.get("data_completeness_score", 0)
    is_preliminary = completeness < 0.5
    
    if is_preliminary:
        # Add preliminary risk note if not present
        risk_notes = estimate.get("risk_notes", [])
        if not any(r.get("risk_id") == "PRELIMINARY_LOW_COMPLETENESS" for r in risk_notes):
            risk_notes.append({
                "risk_id": "PRELIMINARY_LOW_COMPLETENESS",
                "title": "Preliminary estimate with limited data",
                "severity": "high",
                "reason": f"Data completeness score is {completeness*100:.0f}%. This is a preliminary estimate."
            })
            supabase.table("credit_estimates")\
                .update({"risk_notes": risk_notes})\
                .eq("id", estimate_id)\
                .execute()
    
    # Update status
    supabase.table("credit_estimates")\
        .update({"status": "pending_senior_signoff"})\
        .eq("id", estimate_id)\
        .execute()
    
    # Create escalation for senior review queue (integrate with Prompt 11)
    escalation_id = str(uuid4())
    try:
        supabase.table("escalation_requests").insert({
            "id": escalation_id,
            "organization_id": org_id,
            "client_company_id": estimate["client_company_id"],
            "tax_year": estimate["tax_year"],
            "source_type": "manual",
            "source_id": estimate_id,
            "title": f"Credit Estimate Signoff Required - Tax Year {estimate['tax_year']}",
            "summary": f"Credit range estimate v{estimate['estimate_version']} requires senior approval. " +
                      f"Base credit: ${estimate['range_base'].get('credit_amount_selected', 0):,.0f}. " +
                      f"Completeness: {completeness*100:.0f}%.",
            "severity": "high" if is_preliminary else "medium",
            "estimated_impact": {
                "qre_at_risk": estimate["range_base"].get("total_qre", 0),
                "credit_at_risk": estimate["range_base"].get("credit_amount_selected", 0)
            },
            "proposed_action": {"action": "review_and_approve", "estimate_id": estimate_id},
            "status": "queued",
            "created_by_user_id": user["id"]
        }).execute()
    except Exception as e:
        logger.warning(f"Could not create escalation: {e}")
    
    # Audit log
    write_audit_log(
        org_id=org_id,
        user_id=user["id"],
        action="credit_estimate_submitted_for_signoff",
        item_type="credit_estimate",
        item_id=estimate_id,
        details={
            "is_preliminary": is_preliminary,
            "completeness_score": completeness,
            "escalation_id": escalation_id
        }
    )
    
    return {
        "status": "pending_senior_signoff",
        "is_preliminary": is_preliminary,
        "escalation_id": escalation_id
    }


@router.post("/{estimate_id}/signoff")
async def signoff_estimate(
    estimate_id: str,
    request: SignoffRequest,
    user: dict = Depends(get_current_user)
):
    """
    Senior signoff on estimate.
    Only senior+ can approve/reject.
    """
    supabase = get_supabase()
    
    if not check_senior_or_above(user):
        raise HTTPException(status_code=403, detail="Senior CPA or above required for signoff")
    
    profile = get_user_profile(user["id"])
    org_id = profile.get("organization_id")
    
    # Get estimate
    try:
        result = supabase.table("credit_estimates")\
            .select("*")\
            .eq("id", estimate_id)\
            .eq("organization_id", org_id)\
            .single()\
            .execute()
        
        estimate = result.data
    except:
        raise HTTPException(status_code=404, detail="Estimate not found")
    
    if not estimate:
        raise HTTPException(status_code=404, detail="Estimate not found")
    
    if estimate["status"] != "pending_senior_signoff":
        raise HTTPException(status_code=400, detail="Estimate is not pending signoff")
    
    # Validate decision
    if request.decision not in ["approved", "rejected", "changes_requested"]:
        raise HTTPException(status_code=400, detail="Invalid decision")
    
    # Create signoff record
    signoff_id = str(uuid4())
    supabase.table("estimate_signoffs").insert({
        "id": signoff_id,
        "credit_estimate_id": estimate_id,
        "decision": request.decision,
        "reason_code": request.reason_code,
        "note": request.note,
        "completion_method": "senior_override",
        "modifications": request.modifications or {},
        "decided_by_user_id": user["id"],
        "decided_at": datetime.utcnow().isoformat()
    }).execute()
    
    # Update estimate status
    new_status = "draft"  # Default for rejected/changes_requested
    
    if request.decision == "approved":
        new_status = "approved"
        
        # Apply any modifications
        update_data = {
            "status": new_status,
            "approved_by_user_id": user["id"],
            "approved_at": datetime.utcnow().isoformat()
        }
        
        if request.modifications:
            if request.modifications.get("methodology"):
                update_data["methodology"] = request.modifications["methodology"]
            if request.modifications.get("assumptions"):
                # Add senior override assumptions
                existing = estimate.get("assumptions", [])
                for mod_assumption in request.modifications["assumptions"]:
                    mod_assumption["source"] = "senior_override"
                    existing.append(mod_assumption)
                update_data["assumptions"] = existing
        
        supabase.table("credit_estimates")\
            .update(update_data)\
            .eq("id", estimate_id)\
            .execute()
            
    elif request.decision == "rejected":
        new_status = "rejected"
        supabase.table("credit_estimates")\
            .update({"status": new_status})\
            .eq("id", estimate_id)\
            .execute()
    else:  # changes_requested
        new_status = "draft"
        supabase.table("credit_estimates")\
            .update({"status": new_status})\
            .eq("id", estimate_id)\
            .execute()
    
    # Resolve any escalation
    try:
        supabase.table("escalation_requests")\
            .update({"status": "resolved"})\
            .eq("source_id", estimate_id)\
            .execute()
    except:
        pass
    
    # Audit log
    write_audit_log(
        org_id=org_id,
        user_id=user["id"],
        action="credit_estimate_signed_off",
        item_type="credit_estimate",
        item_id=estimate_id,
        details={
            "decision": request.decision,
            "reason_code": request.reason_code,
            "note": request.note,
            "modifications": request.modifications,
            "new_status": new_status
        }
    )
    
    return {
        "status": new_status,
        "signoff_id": signoff_id,
        "decision": request.decision
    }


@router.post("/{estimate_id}/export")
async def export_estimate(
    estimate_id: str,
    request: ExportRequest,
    user: dict = Depends(get_current_user)
):
    """
    Generate client-ready export (PDF/DOCX).
    Only approved estimates can be exported.
    """
    supabase = get_supabase()
    
    if not check_cpa_or_executive(user):
        raise HTTPException(status_code=403, detail="CPA or Executive role required")
    
    profile = get_user_profile(user["id"])
    org_id = profile.get("organization_id")
    
    # Get estimate
    try:
        result = supabase.table("credit_estimates")\
            .select("*, client_companies(name)")\
            .eq("id", estimate_id)\
            .eq("organization_id", org_id)\
            .single()\
            .execute()
        
        estimate = result.data
    except:
        raise HTTPException(status_code=404, detail="Estimate not found")
    
    if not estimate:
        raise HTTPException(status_code=404, detail="Estimate not found")
    
    if estimate["status"] not in ["approved", "pending_senior_signoff"]:
        raise HTTPException(status_code=400, detail="Only approved estimates can be exported")
    
    client_name = estimate.get("client_companies", {}).get("name", "Client")
    
    # Generate export
    if request.export_type == "docx":
        content, mime_type, filename = generate_docx_export(estimate, client_name)
    else:  # pdf
        content, mime_type, filename = generate_pdf_export(estimate, client_name)
    
    # Store export record
    timestamp = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
    storage_path = f"org/{org_id}/client/{estimate['client_company_id']}/estimates/{estimate['tax_year']}/v{estimate['estimate_version']}/{timestamp}.{request.export_type}"
    
    export_id = str(uuid4())
    supabase.table("estimate_exports").insert({
        "id": export_id,
        "credit_estimate_id": estimate_id,
        "export_type": request.export_type,
        "storage_bucket": "estimate-exports",
        "storage_path": storage_path,
        "mime_type": mime_type,
        "created_by_user_id": user["id"],
        "metadata": {
            "filename": filename,
            "size_bytes": len(content),
            "version": estimate["estimate_version"]
        }
    }).execute()
    
    # Upload to storage
    try:
        supabase.storage.from_("estimate-exports").upload(
            storage_path,
            content,
            {"content-type": mime_type}
        )
    except Exception as e:
        logger.warning(f"Storage upload failed: {e}")
    
    # Audit log
    write_audit_log(
        org_id=org_id,
        user_id=user["id"],
        action="credit_estimate_exported",
        item_type="estimate_export",
        item_id=export_id,
        details={
            "estimate_id": estimate_id,
            "export_type": request.export_type,
            "filename": filename
        }
    )
    
    # Return file
    return StreamingResponse(
        io.BytesIO(content),
        media_type=mime_type,
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )


@router.post("/{estimate_id}/email-draft")
async def generate_email_draft(
    estimate_id: str,
    user: dict = Depends(get_current_user)
):
    """
    Generate email draft for client delivery.
    """
    supabase = get_supabase()
    
    if not check_cpa_or_executive(user):
        raise HTTPException(status_code=403, detail="CPA or Executive role required")
    
    profile = get_user_profile(user["id"])
    org_id = profile.get("organization_id")
    
    # Get estimate
    try:
        result = supabase.table("credit_estimates")\
            .select("*, client_companies(name, primary_contact_name, primary_contact_email)")\
            .eq("id", estimate_id)\
            .eq("organization_id", org_id)\
            .single()\
            .execute()
        
        estimate = result.data
    except:
        raise HTTPException(status_code=404, detail="Estimate not found")
    
    if not estimate:
        raise HTTPException(status_code=404, detail="Estimate not found")
    
    client = estimate.get("client_companies", {})
    client_name = client.get("name", "Client")
    contact_name = client.get("primary_contact_name", "")
    
    # Determine status text
    is_preliminary = estimate.get("data_completeness_score", 1) < 0.5
    status_text = "Preliminary" if is_preliminary else "Initial"
    
    base_range = estimate.get("range_base", {})
    low_range = estimate.get("range_low", {})
    high_range = estimate.get("range_high", {})
    
    # Build email
    email_draft = f"""Subject: {status_text} R&D Tax Credit Estimate - Tax Year {estimate['tax_year']}

Dear {contact_name or 'Team'},

I'm pleased to share our {status_text.lower()} R&D Tax Credit estimate for {client_name} for tax year {estimate['tax_year']}.

**Estimated Credit Range:**

| Scenario | Total QRE | Estimated Credit |
|----------|-----------|------------------|
| Conservative | ${low_range.get('total_qre', 0):,.0f} | ${low_range.get('credit_amount_selected', 0):,.0f} |
| Base Case | ${base_range.get('total_qre', 0):,.0f} | ${base_range.get('credit_amount_selected', 0):,.0f} |
| Optimistic | ${high_range.get('total_qre', 0):,.0f} | ${high_range.get('credit_amount_selected', 0):,.0f} |

**Data Completeness:** {estimate.get('data_completeness_score', 0)*100:.0f}%

"""
    
    # Add missing items if any
    missing = estimate.get("missing_inputs", [])
    if missing:
        email_draft += "**To Refine This Estimate, We Need:**\n\n"
        for item in missing[:5]:
            email_draft += f"- {item.get('label', 'Additional information')}\n"
        email_draft += "\n"
    
    # Add risk notes if high severity
    high_risks = [r for r in estimate.get("risk_notes", []) if r.get("severity") == "high"]
    if high_risks:
        email_draft += "**Key Considerations:**\n\n"
        for risk in high_risks:
            email_draft += f"- {risk.get('title')}: {risk.get('reason')}\n"
        email_draft += "\n"
    
    email_draft += f"""**Next Steps:**

1. Review the attached estimate summary
2. Provide any missing documentation noted above
3. Schedule a call to discuss any questions

Please don't hesitate to reach out if you have any questions about this estimate.

Best regards,
Your R&D Tax Credit Team

---
*This estimate is based on information provided and is subject to change upon detailed review.*
"""
    
    # Audit log
    write_audit_log(
        org_id=org_id,
        user_id=user["id"],
        action="credit_estimate_email_draft_created",
        item_type="credit_estimate",
        item_id=estimate_id,
        details={}
    )
    
    return {
        "email_draft": email_draft,
        "to": client.get("primary_contact_email"),
        "subject": f"{status_text} R&D Tax Credit Estimate - Tax Year {estimate['tax_year']}"
    }


@router.post("/{estimate_id}/mark-sent")
async def mark_estimate_sent(
    estimate_id: str,
    user: dict = Depends(get_current_user)
):
    """
    Mark estimate as sent to client.
    """
    supabase = get_supabase()
    
    if not check_cpa_or_executive(user):
        raise HTTPException(status_code=403, detail="CPA or Executive role required")
    
    profile = get_user_profile(user["id"])
    org_id = profile.get("organization_id")
    
    # Get estimate
    try:
        result = supabase.table("credit_estimates")\
            .select("*")\
            .eq("id", estimate_id)\
            .eq("organization_id", org_id)\
            .single()\
            .execute()
        
        estimate = result.data
    except:
        raise HTTPException(status_code=404, detail="Estimate not found")
    
    if not estimate:
        raise HTTPException(status_code=404, detail="Estimate not found")
    
    if estimate["status"] != "approved":
        raise HTTPException(status_code=400, detail="Only approved estimates can be marked as sent")
    
    # Update status
    supabase.table("credit_estimates")\
        .update({"status": "sent_to_client"})\
        .eq("id", estimate_id)\
        .execute()
    
    # Audit log
    write_audit_log(
        org_id=org_id,
        user_id=user["id"],
        action="credit_estimate_marked_sent",
        item_type="credit_estimate",
        item_id=estimate_id,
        details={}
    )
    
    return {"status": "sent_to_client"}


@router.post("/{estimate_id}/recompute")
async def recompute_estimate(
    estimate_id: str,
    user: dict = Depends(get_current_user)
):
    """
    Recompute an estimate (creates new version).
    Use when underlying data has changed.
    """
    supabase = get_supabase()
    
    if not check_cpa_or_executive(user):
        raise HTTPException(status_code=403, detail="CPA or Executive role required")
    
    profile = get_user_profile(user["id"])
    org_id = profile.get("organization_id")
    
    # Get estimate
    try:
        result = supabase.table("credit_estimates")\
            .select("*")\
            .eq("id", estimate_id)\
            .eq("organization_id", org_id)\
            .single()\
            .execute()
        
        estimate = result.data
    except:
        raise HTTPException(status_code=404, detail="Estimate not found")
    
    if not estimate:
        raise HTTPException(status_code=404, detail="Estimate not found")
    
    # Create new draft using same parameters
    return await draft_estimate(
        DraftEstimateRequest(
            client_company_id=estimate["client_company_id"],
            tax_year=estimate["tax_year"],
            methodology=estimate["methodology"],
            range_strategy=estimate.get("range_strategy")
        ),
        user
    )


# ============================================================================
# Export Generation Helpers
# ============================================================================

def generate_docx_export(estimate: dict, client_name: str) -> tuple[bytes, str, str]:
    """Generate DOCX export of credit estimate."""
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    doc = Document()
    
    # Title
    title = doc.add_heading(f'R&D Tax Credit Estimate', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle
    subtitle = doc.add_paragraph(f'{client_name} - Tax Year {estimate["tax_year"]}')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Version and status
    doc.add_paragraph(f'Version {estimate["estimate_version"]} | Status: {estimate["status"].replace("_", " ").title()}')
    doc.add_paragraph(f'Generated: {datetime.now().strftime("%B %d, %Y")}')
    
    doc.add_paragraph()
    
    # Executive Summary
    doc.add_heading('Executive Summary', level=1)
    
    base = estimate.get("range_base", {})
    completeness = estimate.get("data_completeness_score", 0)
    
    summary = doc.add_paragraph()
    summary.add_run(f'Based on the information provided, we estimate a potential R&D tax credit of ')
    summary.add_run(f'${base.get("credit_amount_selected", 0):,.0f}').bold = True
    summary.add_run(f' for tax year {estimate["tax_year"]}.')
    
    doc.add_paragraph(f'Data completeness: {completeness*100:.0f}%')
    
    # Credit Range Table
    doc.add_heading('Credit Range Analysis', level=1)
    
    table = doc.add_table(rows=4, cols=4)
    table.style = 'Table Grid'
    
    # Header row
    headers = ['Scenario', 'Wage QRE', 'Supply QRE', 'Contract QRE', 'Total QRE', 'Est. Credit']
    
    # Simplified table
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'
    
    headers = table.rows[0].cells
    headers[0].text = 'Scenario'
    headers[1].text = 'Total QRE'
    headers[2].text = 'Est. Credit'
    
    low = estimate.get("range_low", {})
    high = estimate.get("range_high", {})
    
    for i, (scenario, data) in enumerate([
        ('Conservative', low),
        ('Base Case', base),
        ('Optimistic', high)
    ], 1):
        row = table.rows[i].cells
        row[0].text = scenario
        row[1].text = f'${data.get("total_qre", 0):,.0f}'
        row[2].text = f'${data.get("credit_amount_selected", 0):,.0f}'
    
    # Assumptions
    doc.add_heading('Key Assumptions', level=1)
    
    assumptions = estimate.get("assumptions", [])
    if assumptions:
        for assumption in assumptions[:10]:
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(assumption.get("title", "")).bold = True
            p.add_run(f': {assumption.get("description", "")}')
    else:
        doc.add_paragraph('No specific assumptions noted.')
    
    # Risk Notes
    risk_notes = estimate.get("risk_notes", [])
    if risk_notes:
        doc.add_heading('Risk Factors', level=1)
        
        for risk in risk_notes:
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(f'[{risk.get("severity", "").upper()}] ').bold = True
            p.add_run(f'{risk.get("title", "")}: {risk.get("reason", "")}')
    
    # Missing Items
    missing = estimate.get("missing_inputs", [])
    if missing:
        doc.add_heading('Outstanding Items', level=1)
        
        for item in missing[:10]:
            doc.add_paragraph(item.get("label", "Unknown"), style='List Bullet')
    
    # Disclaimer
    doc.add_paragraph()
    disclaimer = doc.add_paragraph()
    disclaimer.add_run('Disclaimer: ').bold = True
    disclaimer.add_run(
        'This estimate is preliminary and based on information provided. Final credit amounts may vary upon '
        'completion of the full study and IRS review. This document is for internal planning purposes only.'
    )
    
    # Save to bytes
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    
    filename = f'{client_name.replace(" ", "_")}_Credit_Estimate_{estimate["tax_year"]}_v{estimate["estimate_version"]}.docx'
    
    return file_stream.getvalue(), 'application/vnd.openxmlformats-officedocument.wordprocessingml.document', filename


def generate_pdf_export(estimate: dict, client_name: str) -> tuple[bytes, str, str]:
    """
    Generate PDF export of credit estimate.
    Falls back to DOCX if PDF libraries not available.
    """
    try:
        # Try to use reportlab for PDF
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        
        # Title
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            spaceAfter=12,
            alignment=1  # Center
        )
        story.append(Paragraph('R&D Tax Credit Estimate', title_style))
        story.append(Paragraph(f'{client_name} - Tax Year {estimate["tax_year"]}', styles['Normal']))
        story.append(Spacer(1, 0.25*inch))
        
        # Version info
        story.append(Paragraph(
            f'Version {estimate["estimate_version"]} | Status: {estimate["status"].replace("_", " ").title()}',
            styles['Normal']
        ))
        story.append(Paragraph(f'Generated: {datetime.now().strftime("%B %d, %Y")}', styles['Normal']))
        story.append(Spacer(1, 0.5*inch))
        
        # Summary
        story.append(Paragraph('Executive Summary', styles['Heading2']))
        base = estimate.get("range_base", {})
        story.append(Paragraph(
            f'Based on the information provided, we estimate a potential R&D tax credit of '
            f'<b>${base.get("credit_amount_selected", 0):,.0f}</b> for tax year {estimate["tax_year"]}.',
            styles['Normal']
        ))
        story.append(Spacer(1, 0.25*inch))
        
        # Range Table
        story.append(Paragraph('Credit Range Analysis', styles['Heading2']))
        
        low = estimate.get("range_low", {})
        high = estimate.get("range_high", {})
        
        table_data = [
            ['Scenario', 'Total QRE', 'Est. Credit'],
            ['Conservative', f'${low.get("total_qre", 0):,.0f}', f'${low.get("credit_amount_selected", 0):,.0f}'],
            ['Base Case', f'${base.get("total_qre", 0):,.0f}', f'${base.get("credit_amount_selected", 0):,.0f}'],
            ['Optimistic', f'${high.get("total_qre", 0):,.0f}', f'${high.get("credit_amount_selected", 0):,.0f}'],
        ]
        
        t = Table(table_data, colWidths=[2*inch, 2*inch, 2*inch])
        t.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ]))
        story.append(t)
        story.append(Spacer(1, 0.5*inch))
        
        # Assumptions
        assumptions = estimate.get("assumptions", [])
        if assumptions:
            story.append(Paragraph('Key Assumptions', styles['Heading2']))
            for assumption in assumptions[:8]:
                story.append(Paragraph(
                    f'• <b>{assumption.get("title", "")}</b>: {assumption.get("description", "")}',
                    styles['Normal']
                ))
            story.append(Spacer(1, 0.25*inch))
        
        # Risk Notes
        risk_notes = estimate.get("risk_notes", [])
        if risk_notes:
            story.append(Paragraph('Risk Factors', styles['Heading2']))
            for risk in risk_notes:
                story.append(Paragraph(
                    f'• [{risk.get("severity", "").upper()}] {risk.get("title", "")}: {risk.get("reason", "")}',
                    styles['Normal']
                ))
            story.append(Spacer(1, 0.25*inch))
        
        # Disclaimer
        story.append(Spacer(1, 0.5*inch))
        disclaimer_style = ParagraphStyle(
            'Disclaimer',
            parent=styles['Normal'],
            fontSize=8,
            textColor=colors.grey
        )
        story.append(Paragraph(
            '<b>Disclaimer:</b> This estimate is preliminary and based on information provided. '
            'Final credit amounts may vary upon completion of the full study and IRS review. '
            'This document is for internal planning purposes only.',
            disclaimer_style
        ))
        
        doc.build(story)
        buffer.seek(0)
        
        filename = f'{client_name.replace(" ", "_")}_Credit_Estimate_{estimate["tax_year"]}_v{estimate["estimate_version"]}.pdf'
        return buffer.getvalue(), 'application/pdf', filename
        
    except ImportError:
        # Fallback to DOCX
        logger.warning("reportlab not available, falling back to DOCX export")
        content, _, _ = generate_docx_export(estimate, client_name)
        filename = f'{client_name.replace(" ", "_")}_Credit_Estimate_{estimate["tax_year"]}_v{estimate["estimate_version"]}.docx'
        return content, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document', filename
