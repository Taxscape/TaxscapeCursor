"""
Study Packaging Service (v2)

Implements:
- Readiness evaluation for finalization
- Finalization orchestration
- Artifact generation (Excel, Form 6765 export, narratives, cover summary, ZIP package)
- Signed URL generation

Design notes:
- Uses existing workspace->RDAnalysisSession adapter + Excel generator from app.study_routes/app.rd_excel_generator
- Produces immutable, versioned artifacts under studies_v2
"""

from __future__ import annotations

import csv
import io
import json
import hashlib
import logging
import zipfile
from datetime import datetime, timedelta
from typing import Any, Dict, List, Optional, Tuple
from uuid import uuid4

from openpyxl import load_workbook, Workbook

logger = logging.getLogger(__name__)


STUDY_ARTIFACT_BUCKET = "study-artifacts"


def check_study_locked(supabase, client_id: str, tax_year: int) -> Optional[Dict[str, Any]]:
    """
    Check if a study is locked for the given client/year.
    Returns lock info if locked, None if not locked.
    Used as a guard before canonical table mutations.
    """
    try:
        result = supabase.table("studies_v2").select("id, study_version, status, locked_at, lock_reason").eq("client_company_id", client_id).eq("tax_year", tax_year).in_("status", ["final", "complete"]).order("study_version", desc=True).limit(1).execute()
        if result.data and result.data[0].get("locked_at"):
            return result.data[0]
    except Exception as e:
        logger.warning(f"check_study_locked error: {e}")
    return None

REQUIRED_ARTIFACTS_BASE = [
    "excel_study_workbook",
    "form_6765_export",
    "section_41_narratives_docx",
    "client_cover_summary_pdf",
    "client_package_zip",
]


def _sha256_bytes(b: bytes) -> str:
    return hashlib.sha256(b).hexdigest()


def _safe_filename(name: str) -> str:
    return "".join(c if c.isalnum() or c in ("-", "_", ".", " ") else "_" for c in (name or "file")).strip().replace(" ", "_")


def _now_iso() -> str:
    return datetime.utcnow().isoformat()


def _get_authority_by_keys(supabase, citation_keys: List[str]) -> List[Dict[str, Any]]:
    if not citation_keys:
        return []
    try:
        res = supabase.table("authority_library").select("id,citation_key,citation_label,summary,url").in_("citation_key", citation_keys).execute()
        return res.data or []
    except Exception as e:
        logger.warning(f"authority lookup failed: {e}")
        return []


def _get_latest_intake_session(supabase, client_id: str) -> Optional[Dict[str, Any]]:
    try:
        res = supabase.table("client_intake_sessions").select("*").eq("client_company_id", client_id).order("created_at", desc=True).limit(1).execute()
        return (res.data or [None])[0]
    except Exception:
        return None


def _get_latest_approved_credit_estimate(supabase, client_id: str, tax_year: int) -> Optional[Dict[str, Any]]:
    try:
        res = supabase.table("credit_estimates").select("*").eq("client_company_id", client_id).eq("tax_year", tax_year).eq("status", "approved").order("estimate_version", desc=True).limit(1).execute()
        return (res.data or [None])[0]
    except Exception:
        return None


def evaluate_study_readiness(supabase, org_id: str, client_id: str, tax_year: int, user_id: str) -> Dict[str, Any]:
    """
    Deterministic readiness evaluator.
    Returns checks list + blocker counts and persists to study_finalization_checks.
    """
    checks: List[Dict[str, Any]] = []

    def add(check_id: str, status: str, blocking: bool, message: str, remediation: Optional[Dict[str, Any]] = None, severity: str = "info"):
        checks.append(
            {
                "check_id": check_id,
                "status": status,  # pass|fail|warn
                "blocking": blocking,
                "severity": severity,  # info|warning|error
                "message": message,
                "remediation": remediation or {},
            }
        )

    # 1) Intake session status
    intake = _get_latest_intake_session(supabase, client_id)
    if not intake:
        add(
            "intake_session_exists",
            "fail",
            True,
            "No intake session found for this client.",
            remediation={"target": "Intake Inbox", "href": "/portal/intake-inbox"},
            severity="error",
        )
    else:
        status = intake.get("status")
        if status != "complete":
            add(
                "intake_session_complete",
                "fail",
                True,
                f"Latest intake session is not complete (status={status}).",
                remediation={"target": "Intake Inbox", "href": "/portal/intake-inbox"},
                severity="error",
            )
        else:
            add("intake_session_complete", "pass", False, "Intake session is complete.")

        # expected_inputs completeness (best-effort)
        expected_inputs = intake.get("expected_inputs") or {}
        missing_expected = []
        if isinstance(expected_inputs, dict):
            for k, v in expected_inputs.items():
                if isinstance(v, dict) and not v.get("received", False):
                    missing_expected.append({"key": k, "label": v.get("label", k)})
        if missing_expected:
            add(
                "expected_inputs_received",
                "fail",
                True,
                f"{len(missing_expected)} expected intake input(s) are missing.",
                remediation={"target": "Intake Inbox", "href": "/portal/intake-inbox"},
                severity="error",
            )
        else:
            add("expected_inputs_received", "pass", False, "All expected intake inputs are received.")

    # 2) Review findings
    try:
        findings = supabase.table("review_findings").select("id,severity,status,title").eq("client_company_id", client_id).eq("tax_year", tax_year).execute().data or []
    except Exception:
        findings = []
    open_findings = [f for f in findings if f.get("status") in ("open", "in_review")]
    open_high = [f for f in open_findings if f.get("severity") == "high"]
    open_medium_low = [f for f in open_findings if f.get("severity") in ("medium", "low")]

    if open_high:
        add(
            "no_open_high_findings",
            "fail",
            True,
            f"{len(open_high)} open high-severity finding(s) must be resolved or overridden.",
            remediation={"target": "Review Inbox", "href": "/portal/review"},
            severity="error",
        )
    else:
        add("no_open_high_findings", "pass", False, "No open high-severity findings.")

    if open_medium_low:
        add(
            "open_medium_low_findings_disclosure",
            "warn",
            False,
            f"{len(open_medium_low)} medium/low finding(s) remain; they will be disclosed in the cover summary.",
            remediation={"target": "Review Inbox", "href": "/portal/review"},
            severity="warning",
        )
    else:
        add("open_medium_low_findings_disclosure", "pass", False, "No open medium/low findings.")

    # 3) Escalations
    try:
        escalations = supabase.table("escalation_requests").select("id,status,title").eq("client_company_id", client_id).execute().data or []
    except Exception:
        escalations = []
    open_escalations = [e for e in escalations if e.get("status") in ("queued", "assigned", "in_review", "returned_to_junior")]
    if open_escalations:
        add(
            "no_open_escalations",
            "fail",
            True,
            f"{len(open_escalations)} escalation(s) are still open.",
            remediation={"target": "Senior Queue", "href": "/portal/escalations"},
            severity="error",
        )
    else:
        add("no_open_escalations", "pass", False, "No open escalations.")

    # 4) Evidence requests
    try:
        reqs = supabase.table("evidence_requests").select("id,status,title").eq("client_company_id", client_id).execute().data or []
    except Exception:
        reqs = []
    open_evidence = [r for r in reqs if r.get("status") in ("sent", "awaiting_upload", "partially_received")]
    if open_evidence:
        add(
            "evidence_requests_completed",
            "fail",
            True,
            f"{len(open_evidence)} evidence request(s) are still awaiting completion.",
            remediation={"target": "Evidence Center", "href": "/portal/evidence"},
            severity="error",
        )
    else:
        add("evidence_requests_completed", "pass", False, "No pending evidence requests.")

    # 5) Credit estimate
    approved_est = _get_latest_approved_credit_estimate(supabase, client_id, tax_year)
    if not approved_est:
        add(
            "approved_credit_estimate",
            "fail",
            True,
            "No approved credit estimate exists for this tax year.",
            remediation={"target": "Credit Range", "href": "/portal/credit-range"},
            severity="error",
        )
    else:
        add("approved_credit_estimate", "pass", False, f"Approved credit estimate found (v{approved_est.get('estimate_version')}).")

    # 6) Intake pipeline integrity: pending mappings / parse errors (best-effort)
    try:
        mappings = supabase.table("intake_mappings").select("id,status").eq("client_company_id", client_id).execute().data or []
        pending_mappings = [m for m in mappings if m.get("status") not in ("resolved", "dismissed", "completed")]
    except Exception:
        pending_mappings = []
    if pending_mappings:
        add(
            "no_pending_mappings",
            "fail",
            True,
            f"{len(pending_mappings)} intake mapping task(s) are still pending.",
            remediation={"target": "Intake Inbox", "href": "/portal/intake-inbox"},
            severity="error",
        )
    else:
        add("no_pending_mappings", "pass", False, "No pending mapping tasks.")

    # Summarize counts
    blocking_count = sum(1 for c in checks if c["blocking"] and c["status"] == "fail")
    warning_count = sum(1 for c in checks if c["status"] == "warn")

    # Persist to study_finalization_checks
    try:
        supabase.table("study_finalization_checks").insert(
            {
                "client_company_id": client_id,
                "tax_year": tax_year,
                "computed_at": _now_iso(),
                "checks": checks,
                "blocking_count": blocking_count,
                "warning_count": warning_count,
                "computed_by_user_id": user_id,
            }
        ).execute()
    except Exception as e:
        logger.warning(f"failed to persist study_finalization_checks: {e}")

    return {"checks": checks, "blocking_count": blocking_count, "warning_count": warning_count}


def _snapshot_metadata(supabase, org_id: str, client_id: str, tax_year: int) -> Dict[str, Any]:
    """Capture counts, resolved/dismissed finding IDs, escalation closure summary, evidence summary, and file hashes."""
    snapshot: Dict[str, Any] = {"captured_at": _now_iso(), "generator_version": "study_packaging_v2.1"}

    def count(table: str, filters: List[Tuple[str, Any]]):
        q = supabase.table(table).select("id", count="exact")
        for k, v in filters:
            q = q.eq(k, v)
        r = q.execute()
        return int(getattr(r, "count", None) or len(r.data or []))

    snapshot["counts"] = {
        "employees": count("employees", [("client_company_id", client_id), ("tax_year", tax_year)]),
        "projects": count("projects", [("client_company_id", client_id)]),
        "vendors": count("vendors", [("client_company_id", client_id)]),
        "expenses": count("expenses", [("client_company_id", client_id), ("tax_year", tax_year)]),
        "time_logs": count("time_logs", [("client_company_id", client_id), ("tax_year", tax_year)]),
        "supplies": count("supplies", [("client_company_id", client_id), ("tax_year", tax_year)]),
    }

    # Findings IDs by status
    try:
        f = supabase.table("review_findings").select("id,status,severity").eq("client_company_id", client_id).eq("tax_year", tax_year).execute().data or []
        snapshot["findings"] = {
            "resolved_ids": [x["id"] for x in f if x.get("status", "").startswith("resolved")],
            "dismissed_ids": [x["id"] for x in f if x.get("status") == "dismissed"],
            "open_high_ids": [x["id"] for x in f if x.get("status") in ("open", "in_review") and x.get("severity") == "high"],
        }
    except Exception:
        snapshot["findings"] = {}

    # Escalations summary
    try:
        esc = supabase.table("escalation_requests").select("id,status").eq("client_company_id", client_id).execute().data or []
        snapshot["escalations"] = {
            "closed_ids": [e["id"] for e in esc if e.get("status") in ("resolved", "cancelled")],
            "open_ids": [e["id"] for e in esc if e.get("status") in ("queued", "assigned", "in_review", "returned_to_junior")],
        }
    except Exception:
        snapshot["escalations"] = {}

    # Evidence request summary
    try:
        ev = supabase.table("evidence_requests").select("id,status,request_type").eq("client_company_id", client_id).execute().data or []
        snapshot["evidence_requests"] = {
            "counts_by_status": _counts_by_key(ev, "status"),
            "open_request_ids": [r["id"] for r in ev if r.get("status") in ("sent", "awaiting_upload", "partially_received")],
        }
    except Exception:
        snapshot["evidence_requests"] = {}

    # Input file hashes from intake_files/evidence_files
    try:
        intake_files = supabase.table("intake_files").select("id,sha256").eq("client_company_id", client_id).execute().data or []
        evidence_files = supabase.table("evidence_files").select("id,sha256").eq("client_company_id", client_id).execute().data or []
        snapshot["input_file_hashes"] = {
            "intake_files": [{"id": f["id"], "sha256": f.get("sha256")} for f in intake_files],
            "evidence_files": [{"id": f["id"], "sha256": f.get("sha256")} for f in evidence_files],
        }
    except Exception:
        snapshot["input_file_hashes"] = {}

    return snapshot


def _counts_by_key(rows: List[Dict[str, Any]], key: str) -> Dict[str, int]:
    out: Dict[str, int] = {}
    for r in rows:
        v = str(r.get(key) or "unknown")
        out[v] = out.get(v, 0) + 1
    return out


def _storage_base_path(org_id: str, client_id: str, tax_year: int, version: int) -> str:
    return f"org/{org_id}/client/{client_id}/studies/{tax_year}/v{version}"


def _upsert_artifact_row(
    supabase,
    study_id: str,
    artifact_type: str,
    *,
    status: str,
    storage_bucket: str,
    storage_path: str,
    mime_type: str,
    sha256: str,
    page_count: Optional[int],
    metadata: Dict[str, Any],
    created_by_user_id: str,
    error: Optional[str] = None,
    started_at: Optional[str] = None,
    completed_at: Optional[str] = None,
):
    supabase.table("study_artifacts_v2").upsert(
        {
            "study_id": study_id,
            "artifact_type": artifact_type,
            "generation_status": status,
            "storage_bucket": storage_bucket,
            "storage_path": storage_path,
            "mime_type": mime_type,
            "sha256": sha256,
            "page_count": page_count,
            "metadata": metadata or {},
            "created_by_user_id": created_by_user_id,
            "error": error,
            "started_at": started_at,
            "completed_at": completed_at,
        },
        on_conflict="study_id,artifact_type",
    ).execute()


def generate_excel_study_workbook(supabase, org_id: str, client_id: str, tax_year: int) -> Tuple[bytes, Dict[str, Any]]:
    """
    Generates the 13-sheet Excel study workbook using existing generator.
    """
    # Reuse the adapter from study_routes
    from app.study_routes import adapt_workspace_to_session, compute_inputs_hash
    from app.rd_excel_generator import generate_rd_workbook
    from app.study_routes import calculate_credit

    client_company = supabase.table("client_companies").select("*").eq("id", client_id).single().execute().data
    projects = (supabase.table("projects").select("*").eq("client_company_id", client_id).execute().data or [])
    employees = (supabase.table("employees").select("*").eq("client_company_id", client_id).eq("tax_year", tax_year).execute().data or [])
    vendors = (supabase.table("vendors").select("*").eq("client_company_id", client_id).execute().data or [])
    contracts = (supabase.table("contracts").select("*").eq("client_company_id", client_id).eq("tax_year", tax_year).execute().data or [])
    ap_transactions = (supabase.table("ap_transactions").select("*").eq("client_company_id", client_id).eq("tax_year", tax_year).execute().data or [])
    supplies = (supabase.table("supplies").select("*").eq("client_company_id", client_id).eq("tax_year", tax_year).execute().data or [])

    # Latest evaluations per project
    evaluations: Dict[str, Dict[str, Any]] = {}
    evaluation_ids: List[str] = []
    for p in projects:
        pid = p.get("id")
        if not pid:
            continue
        ev = supabase.table("project_ai_evaluations").select("*").eq("project_id", pid).eq("tax_year", tax_year).order("evaluation_version", desc=True).limit(1).execute().data
        if ev:
            evaluations[pid] = ev[0]
            evaluation_ids.append(ev[0]["id"])

    gaps = (supabase.table("project_gaps").select("*").eq("client_company_id", client_id).eq("tax_year", tax_year).in_("status", ["open", "in_progress"]).execute().data or [])
    qre_summary = (supabase.table("qre_summaries").select("*").eq("client_company_id", client_id).eq("tax_year", tax_year).order("created_at", desc=True).limit(1).execute().data or [None])[0]

    # Adapt
    session = adapt_workspace_to_session(
        client_company=client_company,
        projects=projects,
        employees=employees,
        vendors=vendors,
        contracts=contracts,
        ap_transactions=ap_transactions,
        supplies=supplies,
        evaluations=evaluations,
        gaps=gaps,
        qre_summary=qre_summary,
        tax_year=tax_year,
    )

    # Generate workbook bytes (existing generator returns bytes)
    workbook_bytes = generate_rd_workbook(session)
    inputs_hash = compute_inputs_hash(projects, employees, vendors, ap_transactions + supplies, list(evaluations.values()))
    credit_info = calculate_credit(session.total_qre, "both")

    meta = {
        "inputs_snapshot_hash": inputs_hash,
        "sheet_count": 13,
        "project_count": len(projects),
        "employee_count": len(employees),
        "total_qre": float(session.total_qre),
        "total_credit_recommended": float(credit_info.get("recommended_credit", 0)),
        "evaluation_ids": evaluation_ids,
    }
    return workbook_bytes, meta


def generate_form_6765_export_xlsx(
    *,
    excel_workbook_bytes: bytes,
    approved_credit_estimate: Dict[str, Any],
    client_name: str,
    tax_year: int,
) -> Tuple[bytes, Dict[str, Any]]:
    """
    Produces a client-ready Form 6765 computation artifact as a single-sheet XLSX.
    Uses the approved credit estimate for line amounts where possible.
    """
    wb = load_workbook(io.BytesIO(excel_workbook_bytes))
    sheet_name = "Form_6765_Computation"
    if sheet_name not in wb.sheetnames:
        # Fallback: create a basic sheet
        out = Workbook()
        ws = out.active
        ws.title = "Form_6765"
        ws.append(["Form 6765 computation not available in workbook; generated fallback."])
        b = io.BytesIO()
        out.save(b)
        return b.getvalue(), {"fallback": True}

    ws_in = wb[sheet_name]
    out = Workbook()
    ws_out = out.active
    ws_out.title = "Form_6765"

    # Copy values only (styles omitted for speed)
    max_row = min(ws_in.max_row, 200)
    max_col = min(ws_in.max_column, 10)
    for r in range(1, max_row + 1):
        row_vals = []
        for c in range(1, max_col + 1):
            row_vals.append(ws_in.cell(row=r, column=c).value)
        ws_out.append(row_vals)

    # Add an explicit disclosure + approved estimate reconciliation
    ws_out.append([])
    ws_out.append(["Approved Credit Estimate Reconciliation"])
    base = approved_credit_estimate.get("range_base") or {}
    ws_out.append(["Client", client_name])
    ws_out.append(["Tax Year", tax_year])
    ws_out.append(["QRE (Base)", float(base.get("total_qre") or 0)])
    ws_out.append(["Credit (Regular)", float(base.get("credit_amount_regular") or 0)])
    ws_out.append(["Credit (ASC)", float(base.get("credit_amount_asc") or 0)])
    ws_out.append(["Selected Methodology", approved_credit_estimate.get("methodology")])

    b = io.BytesIO()
    out.save(b)
    meta = {"sheet": sheet_name, "rows_copied": max_row, "reconciled_to_estimate_id": approved_credit_estimate.get("id")}
    return b.getvalue(), meta


def generate_section_41_narratives_docx(
    supabase,
    *,
    org_id: str,
    client_id: str,
    tax_year: int,
    client_name: str,
) -> Tuple[bytes, Dict[str, Any]]:
    """
    Generates Section 41 narratives / four-part test documentation bundle (DOCX).
    Authority refs are pulled from authority_library only.
    """
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    # Authority refs (core)
    authorities = _get_authority_by_keys(supabase, ["IRC_41_D", "IRC_41_B", "FORM_6765_GENERAL"])

    projects = (supabase.table("projects").select("*").eq("client_company_id", client_id).execute().data or [])

    # Latest evaluations per project
    evals: Dict[str, Dict[str, Any]] = {}
    for p in projects:
        pid = p.get("id")
        if not pid:
            continue
        ev = supabase.table("project_ai_evaluations").select("*").eq("project_id", pid).eq("tax_year", tax_year).order("evaluation_version", desc=True).limit(1).execute().data
        if ev:
            evals[pid] = ev[0]

    doc = Document()
    title = doc.add_paragraph("Section 41 Project Narratives (Four-Part Test)")
    title.runs[0].font.size = Pt(18)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"{client_name} — Tax Year {tax_year}").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Generated: {datetime.utcnow().strftime('%Y-%m-%d')}").alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    if authorities:
        doc.add_heading("Authority References (Library-backed)", level=1)
        for a in authorities:
            doc.add_paragraph(f"{a.get('citation_label')} ({a.get('citation_key')})", style="List Bullet")
            doc.add_paragraph(a.get("summary") or "")
        doc.add_page_break()

    qualified_count = 0
    for proj in projects:
        pid = proj.get("id")
        ev = evals.get(pid, {})
        qualified = bool(ev.get("qualified_boolean", False)) if ev else (proj.get("qualification_status") == "qualified")
        if qualified:
            qualified_count += 1

        doc.add_heading(f"Project: {proj.get('name', 'Unnamed')} ({pid})", level=2)
        doc.add_paragraph(f"Status: {'Qualified' if qualified else 'Not Qualified/Excluded'}")
        doc.add_paragraph(f"Category/Product Line: {proj.get('category') or proj.get('product_line') or '—'}")
        if proj.get("description"):
            doc.add_paragraph(f"Description: {proj.get('description')}")

        # Four-part test fields
        doc.add_heading("Four-Part Test", level=3)
        doc.add_paragraph(f"Permitted Purpose: {proj.get('permitted_purpose_type') or proj.get('permitted_purpose') or '—'}")
        doc.add_paragraph(f"Technical Uncertainty: {proj.get('technical_uncertainty') or proj.get('uncertainty_type') or proj.get('uncertainty_statement') or '—'}")
        doc.add_paragraph(f"Process of Experimentation: {proj.get('process_of_experimentation') or proj.get('experimentation_summary') or proj.get('experimentation_description') or '—'}")
        doc.add_paragraph(f"Technological in Nature: {proj.get('technological_basis') or proj.get('technological_nature') or '—'}")

        # AI summary (clearly labeled)
        ai_summary = ev.get("ai_summary") or ev.get("summary") or ""
        if ai_summary:
            doc.add_heading("AI Summary (Reviewed)", level=3)
            doc.add_paragraph(ai_summary)

        # Evidence pointers
        try:
            ev_files = supabase.table("evidence_files").select("id,original_filename,created_at").eq("client_company_id", client_id).eq("entity_type", "project").eq("entity_id", pid).execute().data or []
        except Exception:
            ev_files = []
        doc.add_heading("Evidence Pointers", level=3)
        if ev_files:
            for f in ev_files[:15]:
                doc.add_paragraph(f"- {f.get('original_filename')} (evidence_file_id={f.get('id')})")
        else:
            doc.add_paragraph("No linked evidence files on record.")

        # Exclusion explanation
        if not qualified:
            doc.add_heading("Exclusion Note", level=3)
            doc.add_paragraph("This project is not included in qualified research for this study version based on current determinations.")

        doc.add_page_break()

    b = io.BytesIO()
    doc.save(b)
    return b.getvalue(), {"project_count": len(projects), "qualified_projects": qualified_count, "authority_ids": [a["id"] for a in authorities]}


def generate_section_174_docx_if_applicable(
    supabase,
    *,
    client_id: str,
    client_name: str,
    tax_year: int,
) -> Optional[Tuple[bytes, Dict[str, Any]]]:
    """
    Generate Section 174 narratives if purchased.
    """
    try:
        client = supabase.table("client_companies").select("purchased_sections").eq("id", client_id).single().execute().data
        purchased = client.get("purchased_sections") or {"section_41": True, "section_174": False}
        if not purchased.get("section_174", False):
            return None
    except Exception:
        return None

    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    doc = Document()
    title = doc.add_paragraph("Section 174 Capitalization Summary")
    title.runs[0].font.size = Pt(18)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"{client_name} — Tax Year {tax_year}").alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("Overview", level=1)
    doc.add_paragraph("This document summarizes Section 174 capitalization inputs captured in TaxScape for the engagement scope purchased.")

    # Pull existing responses if present
    try:
        responses = supabase.table("section_174_responses").select("*").eq("client_company_id", client_id).eq("tax_year", tax_year).execute().data or []
    except Exception:
        responses = []

    doc.add_heading("Captured Responses", level=1)
    if responses:
        for r in responses[:50]:
            doc.add_paragraph(json.dumps(r, default=str)[:800])
    else:
        doc.add_paragraph("No Section 174 response records found.")

    b = io.BytesIO()
    doc.save(b)
    return b.getvalue(), {"responses_count": len(responses)}


def generate_cover_summary_pdf(
    supabase,
    *,
    client_id: str,
    client_name: str,
    tax_year: int,
    approved_credit_estimate: Dict[str, Any],
    readiness: Dict[str, Any],
) -> Tuple[bytes, Dict[str, Any]]:
    """
    Generate a 1–2 page PDF cover summary using reportlab.
    """
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=letter, topMargin=0.75 * inch, bottomMargin=0.75 * inch)
    styles = getSampleStyleSheet()
    story: List[Any] = []

    base = approved_credit_estimate.get("range_base") or {}
    low = approved_credit_estimate.get("range_low") or {}
    high = approved_credit_estimate.get("range_high") or {}

    story.append(Paragraph("R&D Credit Study — Client Cover Summary", styles["Title"]))
    story.append(Paragraph(f"{client_name} — Tax Year {tax_year}", styles["Heading2"]))
    story.append(Paragraph(f"Prepared: {datetime.utcnow().strftime('%Y-%m-%d')}", styles["Normal"]))
    story.append(Spacer(1, 0.25 * inch))

    story.append(Paragraph("Executive Summary", styles["Heading2"]))
    story.append(
        Paragraph(
            f"Based on finalized workspace data and senior-approved methodology, the estimated credit is "
            f"<b>${float(base.get('credit_amount_selected') or 0):,.0f}</b> (base case).",
            styles["Normal"],
        )
    )
    story.append(Spacer(1, 0.2 * inch))

    table_data = [
        ["Scenario", "Total QRE", "Estimated Credit"],
        ["Conservative", f"${float(low.get('total_qre') or 0):,.0f}", f"${float(low.get('credit_amount_selected') or 0):,.0f}"],
        ["Base Case", f"${float(base.get('total_qre') or 0):,.0f}", f"${float(base.get('credit_amount_selected') or 0):,.0f}"],
        ["Optimistic", f"${float(high.get('total_qre') or 0):,.0f}", f"${float(high.get('credit_amount_selected') or 0):,.0f}"],
    ]
    t = Table(table_data, colWidths=[2.0 * inch, 2.0 * inch, 2.0 * inch])
    t.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#111827")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                ("BACKGROUND", (0, 1), (-1, -1), colors.whitesmoke),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ]
        )
    )
    story.append(Paragraph("Credit Range", styles["Heading2"]))
    story.append(t)
    story.append(Spacer(1, 0.25 * inch))

    # Remaining limitations (warn checks)
    warnings = [c for c in readiness.get("checks", []) if c.get("status") == "warn"]
    if warnings:
        story.append(Paragraph("Remaining Considerations (Disclosed)", styles["Heading2"]))
        for w in warnings[:10]:
            story.append(Paragraph(f"- {w.get('message')}", styles["Normal"]))
        story.append(Spacer(1, 0.15 * inch))

    story.append(Paragraph("Signoff", styles["Heading2"]))
    story.append(Paragraph("Senior CPA: _______________________   Date: ____________", styles["Normal"]))
    story.append(Spacer(1, 0.2 * inch))
    story.append(
        Paragraph(
            "<b>Disclaimer:</b> This package is prepared for tax credit support and is based on information provided. "
            "Final credit may vary upon IRS examination or additional information.",
            styles["Normal"],
        )
    )

    doc.build(story)
    buf.seek(0)
    return buf.getvalue(), {"page_count": 1}


def generate_client_package_zip(
    supabase,
    *,
    study_row: Dict[str, Any],
    artifacts: Dict[str, Dict[str, Any]],
    include_section_174: bool,
) -> Tuple[bytes, Dict[str, Any]]:
    """
    Create a ZIP containing all client-facing artifacts + optional evidence index CSV.
    """
    # Download existing artifacts from storage and bundle
    zip_buf = io.BytesIO()
    with zipfile.ZipFile(zip_buf, "w", zipfile.ZIP_DEFLATED) as zf:
        def add_from_storage(artifact_type: str, filename: str):
            art = artifacts.get(artifact_type)
            if not art:
                return
            data = supabase.storage.from_(art["storage_bucket"]).download(art["storage_path"])
            zf.writestr(filename, data)

        add_from_storage("client_cover_summary_pdf", "cover_summary.pdf")
        add_from_storage("excel_study_workbook", "final_study.xlsx")
        add_from_storage("form_6765_export", "form_6765_export.xlsx")
        add_from_storage("section_41_narratives_docx", "section_41_narratives.docx")
        if include_section_174:
            add_from_storage("section_174_narratives_docx", "section_174_narratives.docx")

        # Evidence index CSV (best-effort)
        try:
            ev_files = supabase.table("evidence_files").select("id,original_filename,entity_type,entity_id,review_finding_id,created_at").eq("client_company_id", study_row["client_company_id"]).execute().data or []
            out = io.StringIO()
            w = csv.writer(out)
            w.writerow(["evidence_file_id", "filename", "entity_type", "entity_id", "review_finding_id", "created_at"])
            for r in ev_files[:2000]:
                w.writerow([r.get("id"), r.get("original_filename"), r.get("entity_type"), r.get("entity_id"), r.get("review_finding_id"), r.get("created_at")])
            zf.writestr("evidence_index.csv", out.getvalue())
        except Exception:
            pass

    zip_buf.seek(0)
    return zip_buf.getvalue(), {"contains": list(zf.namelist()) if hasattr(zf, "namelist") else []}


def create_signed_download_url(supabase, bucket: str, path: str, expires_seconds: int = 3600) -> Optional[str]:
    try:
        res = supabase.storage.from_(bucket).create_signed_url(path, expires_seconds)
        # supabase-py returns dict with signedURL
        if isinstance(res, dict):
            return res.get("signedURL") or res.get("signedUrl")
        return getattr(res, "signedURL", None) or getattr(res, "signedUrl", None)
    except Exception as e:
        logger.warning(f"signed url failed: {e}")
        return None

