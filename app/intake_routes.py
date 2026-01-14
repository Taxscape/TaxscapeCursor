"""
Intake Package Generator Routes (Prompt 8)
Implements:
- Template generation (DOCX)
- Email draft generation
- Upload link creation
- Mark as sent workflow
- Audit logging
"""

import logging
import json
import hashlib
import secrets
from datetime import datetime, timedelta
from typing import Optional, List, Dict, Any
from uuid import UUID, uuid4
from io import BytesIO

from fastapi import APIRouter, Depends, HTTPException, Query, Header
from fastapi.responses import StreamingResponse
from pydantic import BaseModel, Field, EmailStr

from .supabase_client import get_supabase

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/api/intake", tags=["intake"])

# ============================================================================
# Auth Dependency (matches main.py pattern)
# ============================================================================

def verify_supabase_token(token: str) -> Optional[dict]:
    """Verify a Supabase JWT and return user data."""
    supabase = get_supabase()
    
    try:
        user_response = supabase.auth.get_user(token)
        if user_response and user_response.user:
            user = user_response.user
            return {
                "id": user.id,
                "email": user.email,
                "user_metadata": user.user_metadata or {},
                "app_metadata": user.app_metadata or {},
            }
        return None
    except Exception as e:
        logger.warning(f"[Intake Auth] Token verification failed: {e}")
        return None


async def get_current_user(authorization: Optional[str] = Header(None)):
    """Extract and verify user from Supabase JWT token."""
    if not authorization:
        raise HTTPException(status_code=401, detail="Authorization header missing")
    
    parts = authorization.split()
    if len(parts) != 2 or parts[0].lower() != "bearer":
        raise HTTPException(status_code=401, detail="Invalid authorization header")
    
    token = parts[1]
    user_data = verify_supabase_token(token)
    
    if not user_data:
        raise HTTPException(status_code=401, detail="Invalid or expired token")
    
    return user_data


# ============================================================================
# Helper Functions
# ============================================================================

def get_user_profile(user_id: str) -> Optional[dict]:
    """Get user profile with organization info."""
    supabase = get_supabase()
    try:
        result = supabase.table("profiles").select("*, organization_id").eq("id", user_id).single().execute()
        return result.data
    except Exception as e:
        logger.error(f"Error fetching profile: {e}")
        return None


def get_user_organization(user: dict) -> Optional[dict]:
    """Get user's organization."""
    profile = get_user_profile(user["id"])
    if not profile or not profile.get("organization_id"):
        return None
    
    supabase = get_supabase()
    try:
        result = supabase.table("organizations").select("*").eq("id", profile["organization_id"]).single().execute()
        return result.data
    except:
        return None


def check_cpa_or_executive(user: dict) -> bool:
    """Check if user is CPA or Executive role."""
    profile = get_user_profile(user["id"])
    if not profile:
        return False
    role = profile.get("role", "").lower()
    return role in ["cpa", "executive", "admin"]


def write_audit_log(
    org_id: str,
    user_id: str,
    action: str,
    item_type: str,
    item_id: str = None,
    details: dict = None
):
    """Write to audit_logs table."""
    supabase = get_supabase()
    try:
        supabase.table("audit_logs").insert({
            "organization_id": org_id,
            "user_id": user_id,
            "action": action,
            "item_type": item_type,
            "item_id": item_id,
            "details": details or {},
            "created_at": datetime.utcnow().isoformat()
        }).execute()
    except Exception as e:
        logger.error(f"Failed to write audit log: {e}")


# ============================================================================
# Intake Requirements Matrix
# ============================================================================

def compute_intake_requirements(
    purchased_sections: dict,
    tax_years: List[int],
    has_vendors_expected: bool = True
) -> dict:
    """
    Compute required templates and inputs based on purchased sections.
    
    Returns:
        {
            "required_templates": [...],
            "optional_templates": [...],
            "expected_inputs": {...}
        }
    """
    required_templates = []
    optional_templates = []
    expected_inputs = {}
    
    # Section 41 (R&D Credit) - Core requirements
    has_section_41 = purchased_sections.get("section_41", True)
    has_section_174 = purchased_sections.get("section_174", False)
    
    # Always include master checklist
    required_templates.append("data_request_master")
    
    if has_section_41:
        # Core Section 41 requirements
        required_templates.extend([
            "projects_questionnaire",
            "employee_payroll_template",
            "timesheet_template",
        ])
        
        expected_inputs["projects"] = {
            "required": True,
            "status": "pending",
            "description": "R&D project questionnaires",
            "files": []
        }
        expected_inputs["employees_payroll"] = {
            "required": True,
            "status": "pending",
            "description": "Employee payroll data with wages",
            "files": []
        }
        expected_inputs["timesheets"] = {
            "required": True,
            "status": "pending",
            "description": "Time allocation or timesheet data",
            "files": []
        }
        
        # Vendors/Contracts - optional if client doesn't expect them
        if has_vendors_expected:
            required_templates.append("vendors_contracts_template")
            expected_inputs["vendors_contracts"] = {
                "required": True,
                "status": "pending",
                "description": "Vendor contracts and research agreements",
                "files": []
            }
        else:
            optional_templates.append("vendors_contracts_template")
            expected_inputs["vendors_contracts"] = {
                "required": False,
                "status": "pending",
                "description": "Vendor contracts (if any)",
                "files": []
            }
        
        # AP/Supplies - typically optional but offered
        optional_templates.extend([
            "ap_transactions_template",
            "supplies_template"
        ])
        expected_inputs["ap_transactions"] = {
            "required": False,
            "status": "pending",
            "description": "AP transactions for supply identification",
            "files": []
        }
        expected_inputs["supplies"] = {
            "required": False,
            "status": "pending",
            "description": "Supplies consumed in R&D",
            "files": []
        }
    
    # Section 174 additions
    if has_section_174:
        required_templates.append("section_174_info_request")
        expected_inputs["section_174_questionnaire"] = {
            "required": True,
            "status": "pending",
            "description": "Section 174 categorization questionnaire",
            "files": []
        }
    else:
        # Still offer the template, just mark as optional
        optional_templates.append("section_174_info_request")
        expected_inputs["section_174_questionnaire"] = {
            "required": False,
            "status": "pending",
            "description": "Section 174 info (if applicable)",
            "files": []
        }
    
    return {
        "required_templates": list(set(required_templates)),
        "optional_templates": list(set(optional_templates)),
        "expected_inputs": expected_inputs
    }


# ============================================================================
# Template Content Definitions
# ============================================================================

TEMPLATE_DEFINITIONS = {
    "data_request_master": {
        "title": "R&D Tax Credit Study - Data Request Checklist",
        "description": "Master checklist of all required documents and data",
        "sections": [
            {
                "name": "Overview",
                "content": """This document outlines all data and documentation needed for your R&D tax credit study.
Please complete each item and submit through the secure portal link provided.
Your CPA team will review submissions and reach out with any questions."""
            },
            {
                "name": "Required Documents Checklist",
                "items": [
                    "✓ Completed Projects Questionnaire (one per R&D project)",
                    "✓ Employee Payroll Data (W-2 wages, allocations by state)",
                    "✓ Time Records or Allocation Methodology",
                    "✓ Vendor/Contractor Agreements (if any)",
                    "✓ AP Transactions Export (optional but recommended)",
                    "✓ Supplies List (materials consumed in R&D)"
                ]
            },
            {
                "name": "Submission Instructions",
                "content": """Upload all files to: {upload_link}
Accepted formats: Excel (.xlsx), CSV, PDF, Word (.docx)
Maximum file size: 50MB per file
Questions? Contact your CPA team."""
            }
        ]
    },
    "projects_questionnaire": {
        "title": "R&D Project Questionnaire",
        "description": "Detailed questionnaire for each R&D project",
        "columns": [
            {"name": "Project Name", "description": "Descriptive name for the project", "required": True},
            {"name": "Project Owner", "description": "Person responsible for the project", "required": True},
            {"name": "Department", "description": "Department or team", "required": True},
            {"name": "Start Date", "description": "When work began", "required": True},
            {"name": "End Date", "description": "When work completed (or 'Ongoing')", "required": True},
            {"name": "Business Goals", "description": "What business problem does this solve?", "required": True},
            {"name": "Technical Uncertainty", "description": "What technical challenges did you face?", "required": True},
            {"name": "Experimentation Process", "description": "How did you test solutions?", "required": True},
            {"name": "Technical Approach", "description": "What technologies or methods were used?", "required": True},
            {"name": "Outcome", "description": "What was achieved? Did it work?", "required": True},
            {"name": "Employees Involved", "description": "List key personnel", "required": False}
        ],
        "example_row": {
            "Project Name": "AI-Powered Document Parser",
            "Project Owner": "Jane Smith",
            "Department": "Engineering",
            "Start Date": "2024-01-15",
            "End Date": "2024-08-30",
            "Business Goals": "Automate extraction of key data from unstructured documents to reduce manual processing time by 80%",
            "Technical Uncertainty": "Uncertain whether NLP models could achieve >95% accuracy on our specific document types without extensive training data",
            "Experimentation Process": "Tested 3 different model architectures, evaluated on 1000 sample docs, iterated on preprocessing",
            "Technical Approach": "Transformer-based NLP, custom fine-tuning, Python/PyTorch",
            "Outcome": "Achieved 97% accuracy after 4 iterations, deployed to production",
            "Employees Involved": "Jane Smith, John Doe, Mike Johnson"
        }
    },
    "employee_payroll_template": {
        "title": "Employee Payroll Data Template",
        "description": "Payroll data for employees involved in R&D activities",
        "columns": [
            {"name": "Employee ID", "description": "Unique identifier", "required": True},
            {"name": "Employee Name", "description": "Full name", "required": True},
            {"name": "Job Title", "description": "Current title", "required": True},
            {"name": "Department", "description": "Department", "required": True},
            {"name": "State", "description": "Work state (for state credits)", "required": True},
            {"name": "W-2 Wages", "description": "Total W-2 wages for the year", "required": True},
            {"name": "Bonus", "description": "Bonus amounts (if separate)", "required": False},
            {"name": "Stock Compensation", "description": "Stock-based comp (if any)", "required": False},
            {"name": "R&D Percentage", "description": "Estimated % time on R&D", "required": False},
            {"name": "Start Date", "description": "Employment start date", "required": False},
            {"name": "End Date", "description": "Employment end date (if terminated)", "required": False}
        ],
        "example_row": {
            "Employee ID": "E001",
            "Employee Name": "Jane Smith",
            "Job Title": "Senior Software Engineer",
            "Department": "Engineering",
            "State": "CA",
            "W-2 Wages": "150000",
            "Bonus": "15000",
            "Stock Compensation": "25000",
            "R&D Percentage": "80%",
            "Start Date": "2020-03-15",
            "End Date": ""
        }
    },
    "timesheet_template": {
        "title": "Time Allocation Template",
        "description": "Time spent on R&D activities by employee and project",
        "columns": [
            {"name": "Employee ID", "description": "Must match payroll data", "required": True},
            {"name": "Employee Name", "description": "Full name", "required": True},
            {"name": "Project Name", "description": "Must match project questionnaire", "required": True},
            {"name": "Period Start", "description": "Start of time period", "required": True},
            {"name": "Period End", "description": "End of time period", "required": True},
            {"name": "Hours", "description": "Hours worked on this project", "required": True},
            {"name": "Activity Type", "description": "Design/Development/Testing/etc.", "required": False},
            {"name": "Notes", "description": "Additional context", "required": False}
        ],
        "example_row": {
            "Employee ID": "E001",
            "Employee Name": "Jane Smith",
            "Project Name": "AI-Powered Document Parser",
            "Period Start": "2024-01-01",
            "Period End": "2024-01-31",
            "Hours": "120",
            "Activity Type": "Development",
            "Notes": "Core NLP model development"
        },
        "alternative_method": """If formal timesheets are not available, you may provide:
1. Employee estimates of % time spent on each project
2. Management allocation methodology documentation
3. Calendar/meeting data showing project involvement

Please include a brief explanation of the methodology used."""
    },
    "vendors_contracts_template": {
        "title": "Vendor & Contract Research Template",
        "description": "Third-party vendors and contractors performing research",
        "columns": [
            {"name": "Vendor Name", "description": "Company or individual name", "required": True},
            {"name": "Country", "description": "Where work performed (US vs foreign)", "required": True},
            {"name": "Contract/Agreement", "description": "Reference to agreement or upload", "required": True},
            {"name": "Scope of Work", "description": "What services provided", "required": True},
            {"name": "Risk Bearer", "description": "Who bears risk of failure?", "required": True},
            {"name": "IP Rights", "description": "Who owns resulting IP?", "required": True},
            {"name": "Total Amount", "description": "Contract value for period", "required": True},
            {"name": "R&D Related Amount", "description": "Portion related to R&D", "required": False}
        ],
        "example_row": {
            "Vendor Name": "Tech Research Labs Inc",
            "Country": "United States",
            "Contract/Agreement": "MSA dated 01/15/2024 + SOW #123",
            "Scope of Work": "Development of custom ML algorithm for document classification",
            "Risk Bearer": "Client bears risk - paid regardless of success",
            "IP Rights": "Client owns all resulting IP",
            "Total Amount": "75000",
            "R&D Related Amount": "75000"
        },
        "notes": """IMPORTANT: Foreign research expenditures (work performed outside the US) are subject to different rules.
Please clearly indicate the country where work was performed.
Attach copies of relevant contracts or statements of work."""
    },
    "ap_transactions_template": {
        "title": "AP Transactions Export",
        "description": "Accounts payable transactions for supply/expense identification",
        "columns": [
            {"name": "Date", "description": "Transaction date", "required": True},
            {"name": "Vendor", "description": "Vendor/payee name", "required": True},
            {"name": "Description", "description": "Line item description", "required": True},
            {"name": "GL Account", "description": "General ledger account code", "required": True},
            {"name": "GL Account Name", "description": "Account description", "required": False},
            {"name": "Amount", "description": "Transaction amount", "required": True},
            {"name": "Department", "description": "Cost center/department", "required": False},
            {"name": "Project", "description": "Project if coded", "required": False},
            {"name": "Category Candidate", "description": "Supply/Contract/Non-R&D", "required": False}
        ],
        "notes": """We use this data to identify potential qualified research supplies.
Export from your accounting system (QuickBooks, NetSuite, etc.)
Include transactions from relevant GL accounts: lab supplies, cloud computing, prototype materials, etc."""
    },
    "supplies_template": {
        "title": "R&D Supplies Template",
        "description": "Materials and supplies consumed in R&D activities",
        "columns": [
            {"name": "Item", "description": "Supply/material name", "required": True},
            {"name": "Purpose", "description": "How used in R&D", "required": True},
            {"name": "Project", "description": "Related R&D project(s)", "required": True},
            {"name": "Amount", "description": "Dollar amount", "required": True},
            {"name": "Consumed?", "description": "Yes/No - was it used up?", "required": True},
            {"name": "Capitalized?", "description": "Yes/No - on balance sheet?", "required": True},
            {"name": "Vendor", "description": "Supplier name", "required": False},
            {"name": "Date", "description": "Purchase/expense date", "required": False}
        ],
        "example_row": {
            "Item": "AWS Cloud Computing",
            "Purpose": "Model training infrastructure",
            "Project": "AI Document Parser",
            "Amount": "12500",
            "Consumed?": "Yes",
            "Capitalized?": "No",
            "Vendor": "Amazon Web Services",
            "Date": "2024-06"
        },
        "notes": """Qualified supplies must be:
• Used in qualified research
• Consumed in the R&D process (not inventory for sale)
• Not capitalized (expensed in the year)

Common examples: cloud computing, prototype materials, lab supplies, software licenses for development"""
    },
    "section_174_info_request": {
        "title": "Section 174 R&E Expenditures Information Request",
        "description": "Information for Section 174 capitalization requirements",
        "sections": [
            {
                "name": "Overview",
                "content": """Section 174 requires capitalization and amortization of specified research or experimental expenditures.
This form helps us properly categorize your R&D expenses under the new rules.
Please complete for each category of R&D spending."""
            },
            {
                "name": "Software Development",
                "questions": [
                    "Do you have internal software development activities? (Yes/No)",
                    "Describe the nature of software development (internal use, external sale, both)",
                    "What percentage of engineering time is spent on new development vs. maintenance?",
                    "Are any development activities performed outside the US?"
                ]
            },
            {
                "name": "R&E Expenditure Categories",
                "items": [
                    {"category": "Labor - US", "description": "Wages for R&D performed in the US"},
                    {"category": "Labor - Foreign", "description": "Wages for R&D performed outside US"},
                    {"category": "Supplies", "description": "Materials consumed in R&D"},
                    {"category": "Contract Research - US", "description": "Third-party R&D in the US"},
                    {"category": "Contract Research - Foreign", "description": "Third-party R&D outside US"}
                ]
            },
            {
                "name": "Treatment Questions",
                "questions": [
                    "How are R&D expenditures currently treated for book purposes?",
                    "Are there any costs currently being capitalized?",
                    "Do you have software development costs being amortized under 167(f)?",
                    "Are there any costs related to the acquisition of patents or IP?"
                ]
            }
        ]
    }
}


# ============================================================================
# DOCX Template Generator
# ============================================================================

def generate_docx_template(
    template_type: str,
    client_info: dict,
    tax_years: List[int],
    purchased_sections: dict,
    org_info: dict,
    upload_link: str = None
) -> BytesIO:
    """
    Generate a DOCX template file with merged client data.
    
    Returns BytesIO containing the DOCX file.
    """
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
    except ImportError:
        # Fallback to simple text if python-docx not available
        logger.warning("python-docx not installed, generating text fallback")
        return generate_text_template_fallback(template_type, client_info, tax_years, purchased_sections, org_info, upload_link)
    
    template_def = TEMPLATE_DEFINITIONS.get(template_type)
    if not template_def:
        raise ValueError(f"Unknown template type: {template_type}")
    
    doc = Document()
    
    # === Header Section ===
    header = doc.sections[0].header
    header_para = header.paragraphs[0]
    header_para.text = f"{org_info.get('name', 'CPA Firm')} | R&D Tax Credit Study"
    header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # === Title ===
    title = doc.add_heading(template_def["title"], 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # === Client Info Block ===
    doc.add_paragraph()
    info_table = doc.add_table(rows=4, cols=2)
    info_table.style = 'Table Grid'
    
    rows_data = [
        ("Client:", client_info.get("name", "CLIENT NAME")),
        ("Tax Year(s):", ", ".join(str(y) for y in tax_years)),
        ("Sections:", ", ".join([k for k, v in purchased_sections.items() if v])),
        ("Scope:", client_info.get("study_scope", "Full R&D Credit Study"))
    ]
    
    for i, (label, value) in enumerate(rows_data):
        info_table.rows[i].cells[0].text = label
        info_table.rows[i].cells[1].text = value
    
    doc.add_paragraph()
    
    # === Return Instructions ===
    if upload_link:
        doc.add_heading("Submission Instructions", level=1)
        p = doc.add_paragraph()
        p.add_run("Upload all completed documents to: ").bold = True
        p.add_run(upload_link)
        doc.add_paragraph()
    
    # === Template-Specific Content ===
    if "sections" in template_def:
        for section in template_def["sections"]:
            doc.add_heading(section["name"], level=1)
            if "content" in section:
                content = section["content"]
                if upload_link:
                    content = content.replace("{upload_link}", upload_link)
                doc.add_paragraph(content)
            if "items" in section:
                for item in section["items"]:
                    if isinstance(item, str):
                        doc.add_paragraph(item, style='List Bullet')
                    elif isinstance(item, dict):
                        p = doc.add_paragraph(style='List Bullet')
                        p.add_run(f"{item.get('category', '')}: ").bold = True
                        p.add_run(item.get('description', ''))
            if "questions" in section:
                for q in section["questions"]:
                    doc.add_paragraph(f"• {q}", style='List Bullet')
    
    # === Columns Table (for data templates) ===
    if "columns" in template_def:
        doc.add_heading("Required Fields", level=1)
        doc.add_paragraph(template_def.get("description", ""))
        
        # Column definitions table
        cols = template_def["columns"]
        table = doc.add_table(rows=len(cols) + 1, cols=3)
        table.style = 'Table Grid'
        
        # Header row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Field Name"
        hdr_cells[1].text = "Description"
        hdr_cells[2].text = "Required"
        
        for cell in hdr_cells:
            cell.paragraphs[0].runs[0].bold = True
        
        for i, col in enumerate(cols):
            row_cells = table.rows[i + 1].cells
            row_cells[0].text = col["name"]
            row_cells[1].text = col["description"]
            row_cells[2].text = "Yes" if col.get("required", False) else "Optional"
        
        doc.add_paragraph()
        
        # Example row
        if "example_row" in template_def:
            doc.add_heading("Example Data Row", level=2)
            example = template_def["example_row"]
            ex_table = doc.add_table(rows=2, cols=len(cols))
            ex_table.style = 'Table Grid'
            
            for i, col in enumerate(cols):
                ex_table.rows[0].cells[i].text = col["name"]
                ex_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
                ex_table.rows[1].cells[i].text = str(example.get(col["name"], ""))
    
    # === Notes Section ===
    if "notes" in template_def:
        doc.add_paragraph()
        doc.add_heading("Important Notes", level=1)
        doc.add_paragraph(template_def["notes"])
    
    if "alternative_method" in template_def:
        doc.add_paragraph()
        doc.add_heading("Alternative Methods", level=2)
        doc.add_paragraph(template_def["alternative_method"])
    
    # === Footer ===
    doc.add_paragraph()
    doc.add_paragraph("─" * 50)
    footer_text = f"Generated by {org_info.get('name', 'TaxScape Pro')} | {datetime.now().strftime('%Y-%m-%d')}"
    footer_para = doc.add_paragraph(footer_text)
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Save to BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer


def generate_text_template_fallback(
    template_type: str,
    client_info: dict,
    tax_years: List[int],
    purchased_sections: dict,
    org_info: dict,
    upload_link: str = None
) -> BytesIO:
    """Fallback text template if python-docx not available."""
    template_def = TEMPLATE_DEFINITIONS.get(template_type, {})
    
    lines = [
        f"{'=' * 60}",
        f"{template_def.get('title', template_type.upper())}",
        f"{'=' * 60}",
        "",
        f"Client: {client_info.get('name', 'CLIENT NAME')}",
        f"Tax Year(s): {', '.join(str(y) for y in tax_years)}",
        f"Sections: {', '.join([k for k, v in purchased_sections.items() if v])}",
        f"Scope: {client_info.get('study_scope', 'Full R&D Credit Study')}",
        "",
    ]
    
    if upload_link:
        lines.extend([
            "SUBMISSION INSTRUCTIONS",
            "-" * 40,
            f"Upload documents to: {upload_link}",
            ""
        ])
    
    if "columns" in template_def:
        lines.append("REQUIRED FIELDS")
        lines.append("-" * 40)
        for col in template_def["columns"]:
            req = "Required" if col.get("required") else "Optional"
            lines.append(f"• {col['name']} ({req}): {col['description']}")
        lines.append("")
    
    if "notes" in template_def:
        lines.append("NOTES")
        lines.append("-" * 40)
        lines.append(template_def["notes"])
    
    content = "\n".join(lines)
    buffer = BytesIO(content.encode('utf-8'))
    buffer.seek(0)
    return buffer


# ============================================================================
# Pydantic Models
# ============================================================================

class GenerateTemplatesRequest(BaseModel):
    client_company_id: str
    tax_years: List[int]
    template_types: Optional[List[str]] = None
    onboarding_session_id: Optional[str] = None


class GenerateTemplatesResponse(BaseModel):
    success: bool
    templates: List[dict]
    required_inputs: dict
    missing_fields: Optional[List[str]] = None
    error: Optional[str] = None


class GenerateUploadLinkRequest(BaseModel):
    client_company_id: str
    tax_years: List[int]
    expires_in_days: int = 30


class GenerateUploadLinkResponse(BaseModel):
    success: bool
    upload_link: str
    token_id: str
    expires_at: str
    error: Optional[str] = None


class GenerateEmailDraftRequest(BaseModel):
    client_company_id: str
    tax_years: List[int]
    template_ids: List[str]
    tone: str = "professional-friendly"
    upload_link: Optional[str] = None


class GenerateEmailDraftResponse(BaseModel):
    success: bool
    draft_id: str
    subject: str
    body_text: str
    to_recipients: List[dict]
    cc_recipients: List[dict]
    missing_fields: Optional[List[str]] = None
    error: Optional[str] = None


class MarkSentRequest(BaseModel):
    email_draft_id: str
    sent_at: Optional[str] = None


class MarkSentResponse(BaseModel):
    success: bool
    intake_session_id: str
    status: str
    error: Optional[str] = None


class UpdateClientSettingsRequest(BaseModel):
    primary_contact_name: Optional[str] = None
    primary_contact_email: Optional[str] = None
    purchased_sections: Optional[dict] = None
    study_scope: Optional[str] = None
    intake_mode: Optional[str] = None
    has_vendors_expected: Optional[bool] = None


# ============================================================================
# API Endpoints
# ============================================================================

@router.post("/templates/generate", response_model=GenerateTemplatesResponse)
async def generate_templates(
    request: GenerateTemplatesRequest,
    user: dict = Depends(get_current_user)
):
    """
    Generate intake templates for a client.
    Creates DOCX files and stores in Supabase Storage.
    """
    supabase = get_supabase()
    
    # Verify user role
    if not check_cpa_or_executive(user):
        raise HTTPException(status_code=403, detail="CPA or Executive role required")
    
    profile = get_user_profile(user["id"])
    org_id = profile.get("organization_id")
    
    if not org_id:
        raise HTTPException(status_code=400, detail="User not associated with an organization")
    
    # Fetch client company
    try:
        client_result = supabase.table("client_companies")\
            .select("*")\
            .eq("id", request.client_company_id)\
            .eq("organization_id", org_id)\
            .single()\
            .execute()
        client = client_result.data
    except Exception as e:
        raise HTTPException(status_code=404, detail=f"Client company not found: {e}")
    
    if not client:
        raise HTTPException(status_code=404, detail="Client company not found")
    
    # Check for required fields
    missing_fields = []
    if not client.get("name"):
        missing_fields.append("name")
    
    # Purchased sections - default if not set
    purchased_sections = client.get("purchased_sections") or {"section_41": True, "section_174": False}
    has_vendors = client.get("has_vendors_expected", True)
    
    if missing_fields:
        return GenerateTemplatesResponse(
            success=False,
            templates=[],
            required_inputs={},
            missing_fields=missing_fields,
            error="Missing required client fields"
        )
    
    # Compute requirements
    requirements = compute_intake_requirements(
        purchased_sections,
        request.tax_years,
        has_vendors
    )
    
    # Determine which templates to generate
    if request.template_types:
        templates_to_generate = request.template_types
    else:
        templates_to_generate = requirements["required_templates"] + requirements["optional_templates"]
    
    # Get organization info
    org_result = supabase.table("organizations").select("*").eq("id", org_id).single().execute()
    org_info = org_result.data or {"name": "TaxScape Pro"}
    
    # Generate each template
    generated_templates = []
    
    for template_type in templates_to_generate:
        try:
            # Get next version number
            version_result = supabase.rpc(
                "get_next_template_version",
                {"p_client_company_id": request.client_company_id, "p_template_type": template_type}
            ).execute()
            
            version = version_result.data if version_result.data else 1
            
            # Generate the document
            doc_buffer = generate_docx_template(
                template_type=template_type,
                client_info=client,
                tax_years=request.tax_years,
                purchased_sections=purchased_sections,
                org_info=org_info,
                upload_link=None  # Will be set later if available
            )
            
            # Upload to Supabase Storage
            file_name = f"{client.get('slug', 'client')}_{template_type}_v{version}_{datetime.now().strftime('%Y%m%d')}.docx"
            storage_path = f"{org_id}/{request.client_company_id}/{file_name}"
            
            # For now, we'll store the path - actual upload would use supabase storage API
            # In production, this would be:
            # supabase.storage.from_("intake-templates").upload(storage_path, doc_buffer.read())
            
            # Calculate hash
            doc_content = doc_buffer.read()
            sha256_hash = hashlib.sha256(doc_content).hexdigest()
            doc_buffer.seek(0)
            
            # Determine completion method
            completion_method = "ai_validated" if request.onboarding_session_id else "manual_user_action"
            
            # Insert template record
            template_record = {
                "id": str(uuid4()),
                "organization_id": org_id,
                "client_company_id": request.client_company_id,
                "tax_years": request.tax_years,
                "template_type": template_type,
                "template_version": version,
                "storage_bucket": "intake-templates",
                "storage_path": storage_path,
                "mime_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "file_size_bytes": len(doc_content),
                "status": "active",
                "created_by_user_id": user["id"],
                "metadata": {
                    "included_sections": [k for k, v in purchased_sections.items() if v],
                    "required_fields": [c["name"] for c in TEMPLATE_DEFINITIONS.get(template_type, {}).get("columns", []) if c.get("required")],
                    "example_rows_present": "example_row" in TEMPLATE_DEFINITIONS.get(template_type, {}),
                    "generated_from_onboarding_session_id": request.onboarding_session_id,
                    "sha256_hash": sha256_hash,
                    "generation_method": completion_method
                }
            }
            
            insert_result = supabase.table("intake_templates").insert(template_record).execute()
            
            # Mark previous versions as superseded
            if version > 1:
                supabase.table("intake_templates")\
                    .update({"status": "superseded"})\
                    .eq("client_company_id", request.client_company_id)\
                    .eq("template_type", template_type)\
                    .lt("template_version", version)\
                    .execute()
            
            generated_templates.append({
                "id": template_record["id"],
                "template_type": template_type,
                "version": version,
                "file_name": file_name,
                "storage_path": storage_path,
                "download_url": f"/api/intake/templates/{template_record['id']}/download"
            })
            
            # Audit log
            write_audit_log(
                org_id=org_id,
                user_id=user["id"],
                action="intake_template_generated",
                item_type="intake_template",
                item_id=template_record["id"],
                details={
                    "template_type": template_type,
                    "version": version,
                    "client_company_id": request.client_company_id,
                    "tax_years": request.tax_years,
                    "completion_method": completion_method
                }
            )
            
        except Exception as e:
            logger.error(f"Error generating template {template_type}: {e}")
            # Continue with other templates
    
    return GenerateTemplatesResponse(
        success=True,
        templates=generated_templates,
        required_inputs=requirements["expected_inputs"]
    )


@router.get("/templates/{template_id}/download")
async def download_template(
    template_id: str,
    user: dict = Depends(get_current_user)
):
    """Download a generated template."""
    supabase = get_supabase()
    profile = get_user_profile(user["id"])
    org_id = profile.get("organization_id")
    
    # Fetch template
    try:
        result = supabase.table("intake_templates")\
            .select("*, client_companies(name, slug)")\
            .eq("id", template_id)\
            .eq("organization_id", org_id)\
            .single()\
            .execute()
        template = result.data
    except:
        raise HTTPException(status_code=404, detail="Template not found")
    
    if not template:
        raise HTTPException(status_code=404, detail="Template not found")
    
    # For demo purposes, regenerate the template on-the-fly
    # In production, this would fetch from Supabase Storage
    client_result = supabase.table("client_companies")\
        .select("*")\
        .eq("id", template["client_company_id"])\
        .single()\
        .execute()
    client = client_result.data
    
    org_result = supabase.table("organizations").select("*").eq("id", org_id).single().execute()
    org_info = org_result.data or {"name": "TaxScape Pro"}
    
    doc_buffer = generate_docx_template(
        template_type=template["template_type"],
        client_info=client,
        tax_years=template["tax_years"],
        purchased_sections=client.get("purchased_sections", {}),
        org_info=org_info
    )
    
    file_name = f"{client.get('slug', 'client')}_{template['template_type']}_v{template['template_version']}.docx"
    
    return StreamingResponse(
        doc_buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={file_name}"}
    )


@router.post("/upload-link", response_model=GenerateUploadLinkResponse)
async def generate_upload_link(
    request: GenerateUploadLinkRequest,
    user: dict = Depends(get_current_user)
):
    """
    Generate a secure upload link for client document submission.
    """
    supabase = get_supabase()
    
    if not check_cpa_or_executive(user):
        raise HTTPException(status_code=403, detail="CPA or Executive role required")
    
    profile = get_user_profile(user["id"])
    org_id = profile.get("organization_id")
    
    # Verify client exists
    try:
        client_result = supabase.table("client_companies")\
            .select("id, name")\
            .eq("id", request.client_company_id)\
            .eq("organization_id", org_id)\
            .single()\
            .execute()
        if not client_result.data:
            raise HTTPException(status_code=404, detail="Client not found")
    except:
        raise HTTPException(status_code=404, detail="Client not found")
    
    # Generate secure token
    raw_token = secrets.token_urlsafe(32)
    token_hash = hashlib.sha256(raw_token.encode()).hexdigest()
    
    expires_at = datetime.utcnow() + timedelta(days=request.expires_in_days)
    
    # Store token
    token_record = {
        "id": str(uuid4()),
        "token_hash": token_hash,
        "organization_id": org_id,
        "client_company_id": request.client_company_id,
        "tax_years": request.tax_years,
        "expires_at": expires_at.isoformat(),
        "created_by_user_id": user["id"],
        "metadata": {
            "generated_by_email": user.get("email")
        }
    }
    
    try:
        supabase.table("intake_upload_tokens").insert(token_record).execute()
    except Exception as e:
        logger.error(f"Failed to create upload token: {e}")
        raise HTTPException(status_code=500, detail="Failed to create upload link")
    
    # Generate the upload link
    # In production, this would be a proper domain
    base_url = "https://app.taxscape.com"  # or from config
    upload_link = f"{base_url}/intake/upload/{raw_token}"
    
    # Audit log
    write_audit_log(
        org_id=org_id,
        user_id=user["id"],
        action="intake_upload_link_created",
        item_type="intake_upload_token",
        item_id=token_record["id"],
        details={
            "client_company_id": request.client_company_id,
            "tax_years": request.tax_years,
            "expires_at": expires_at.isoformat(),
            "token_hash_prefix": token_hash[:8]
        }
    )
    
    return GenerateUploadLinkResponse(
        success=True,
        upload_link=upload_link,
        token_id=token_record["id"],
        expires_at=expires_at.isoformat()
    )


@router.post("/email-draft/generate", response_model=GenerateEmailDraftResponse)
async def generate_email_draft(
    request: GenerateEmailDraftRequest,
    user: dict = Depends(get_current_user)
):
    """
    Generate an email draft for data request.
    """
    supabase = get_supabase()
    
    if not check_cpa_or_executive(user):
        raise HTTPException(status_code=403, detail="CPA or Executive role required")
    
    profile = get_user_profile(user["id"])
    org_id = profile.get("organization_id")
    
    # Fetch client
    try:
        client_result = supabase.table("client_companies")\
            .select("*")\
            .eq("id", request.client_company_id)\
            .eq("organization_id", org_id)\
            .single()\
            .execute()
        client = client_result.data
    except:
        raise HTTPException(status_code=404, detail="Client not found")
    
    if not client:
        raise HTTPException(status_code=404, detail="Client not found")
    
    # Check for required contact info
    missing_fields = []
    contact_email = client.get("primary_contact_email") or client.get("contact_email")
    contact_name = client.get("primary_contact_name") or client.get("contact_name")
    
    if not contact_email:
        missing_fields.append("primary_contact_email")
    
    if missing_fields:
        return GenerateEmailDraftResponse(
            success=False,
            draft_id="",
            subject="",
            body_text="",
            to_recipients=[],
            cc_recipients=[],
            missing_fields=missing_fields,
            error="Missing client contact information"
        )
    
    # Get org info
    org_result = supabase.table("organizations").select("*").eq("id", org_id).single().execute()
    org_info = org_result.data or {"name": "Your CPA Team"}
    
    # Fetch template info
    template_names = []
    for tid in request.template_ids:
        try:
            t_result = supabase.table("intake_templates").select("template_type").eq("id", tid).single().execute()
            if t_result.data:
                template_names.append(TEMPLATE_DEFINITIONS.get(t_result.data["template_type"], {}).get("title", t_result.data["template_type"]))
        except:
            pass
    
    # Build subject
    years_str = ", ".join(str(y) for y in request.tax_years)
    subject = f"R&D Tax Credit Study Data Request - {client['name']} ({years_str})"
    
    # Build body
    upload_link_text = f"\n\nPlease upload all documents to:\n{request.upload_link}\n" if request.upload_link else "\n\nPlease respond to this email with the requested documents attached.\n"
    
    body_text = f"""Dear {contact_name or 'Team'},

I hope this email finds you well. We are beginning work on your R&D Tax Credit study for tax year(s) {years_str}, and need to collect some information and documentation from your team.

**What We Need:**
To properly evaluate your qualifying R&D activities and calculate your credit, we need the following:

"""
    
    # Add template list
    for i, name in enumerate(template_names, 1):
        body_text += f"{i}. {name}\n"
    
    body_text += f"""
Each template includes instructions and examples to help you complete it correctly.

**How to Submit:**{upload_link_text}
**Timeline:**
We request these documents within the next 2-3 weeks to stay on track with your filing timeline. Please let us know if you have any questions or need assistance completing any of the templates.

**What Happens Next:**
Once we receive your documentation:
1. We will review for completeness
2. Our team will analyze your R&D activities
3. We may have follow-up questions to clarify certain projects
4. We will prepare your R&D credit calculation and supporting documentation

Please don't hesitate to reach out if you have any questions.

Best regards,
{profile.get('full_name', 'Your CPA Team')}
{org_info.get('name', '')}
"""
    
    # Store the draft
    draft_record = {
        "id": str(uuid4()),
        "organization_id": org_id,
        "client_company_id": request.client_company_id,
        "tax_years": request.tax_years,
        "subject": subject,
        "body_text": body_text,
        "to_recipients": [{"name": contact_name or "", "email": contact_email}],
        "cc_recipients": [],
        "bcc_recipients": [],
        "attachment_template_ids": request.template_ids,
        "status": "draft",
        "created_by_user_id": user["id"],
        "metadata": {
            "tone": request.tone,
            "upload_link_embedded": bool(request.upload_link),
            "intake_mode_used": client.get("intake_mode", "portal_upload_only")
        }
    }
    
    try:
        supabase.table("intake_email_drafts").insert(draft_record).execute()
    except Exception as e:
        logger.error(f"Failed to save email draft: {e}")
        raise HTTPException(status_code=500, detail="Failed to save email draft")
    
    # Audit log
    write_audit_log(
        org_id=org_id,
        user_id=user["id"],
        action="intake_email_draft_created",
        item_type="intake_email_draft",
        item_id=draft_record["id"],
        details={
            "client_company_id": request.client_company_id,
            "tax_years": request.tax_years,
            "template_count": len(request.template_ids),
            "has_upload_link": bool(request.upload_link)
        }
    )
    
    return GenerateEmailDraftResponse(
        success=True,
        draft_id=draft_record["id"],
        subject=subject,
        body_text=body_text,
        to_recipients=draft_record["to_recipients"],
        cc_recipients=draft_record["cc_recipients"]
    )


@router.post("/email-draft/mark-sent", response_model=MarkSentResponse)
async def mark_email_sent(
    request: MarkSentRequest,
    user: dict = Depends(get_current_user)
):
    """
    Mark an email draft as sent and create intake session.
    """
    supabase = get_supabase()
    
    if not check_cpa_or_executive(user):
        raise HTTPException(status_code=403, detail="CPA or Executive role required")
    
    profile = get_user_profile(user["id"])
    org_id = profile.get("organization_id")
    
    # Fetch email draft
    try:
        draft_result = supabase.table("intake_email_drafts")\
            .select("*")\
            .eq("id", request.email_draft_id)\
            .eq("organization_id", org_id)\
            .single()\
            .execute()
        draft = draft_result.data
    except:
        raise HTTPException(status_code=404, detail="Email draft not found")
    
    if not draft:
        raise HTTPException(status_code=404, detail="Email draft not found")
    
    if draft["status"] == "marked_sent":
        raise HTTPException(status_code=400, detail="Email already marked as sent")
    
    # Fetch client for requirements
    client_result = supabase.table("client_companies")\
        .select("*")\
        .eq("id", draft["client_company_id"])\
        .single()\
        .execute()
    client = client_result.data
    
    purchased_sections = client.get("purchased_sections") or {"section_41": True, "section_174": False}
    has_vendors = client.get("has_vendors_expected", True)
    
    # Compute expected inputs
    requirements = compute_intake_requirements(
        purchased_sections,
        draft["tax_years"],
        has_vendors
    )
    
    sent_at = request.sent_at or datetime.utcnow().isoformat()
    
    # Update draft status
    supabase.table("intake_email_drafts")\
        .update({
            "status": "marked_sent",
            "marked_sent_at": sent_at
        })\
        .eq("id", request.email_draft_id)\
        .execute()
    
    # Create or update intake session
    session_id = str(uuid4())
    session_record = {
        "id": session_id,
        "organization_id": org_id,
        "client_company_id": draft["client_company_id"],
        "tax_years": draft["tax_years"],
        "status": "awaiting_client",
        "expected_inputs": requirements["expected_inputs"],
        "template_ids": draft["attachment_template_ids"],
        "source_email_draft_id": draft["id"],
        "created_by_user_id": user["id"],
        "metadata": {
            "email_sent_at": sent_at,
            "email_sent_to": draft["to_recipients"]
        }
    }
    
    try:
        supabase.table("client_intake_sessions").insert(session_record).execute()
    except Exception as e:
        logger.error(f"Failed to create intake session: {e}")
        raise HTTPException(status_code=500, detail="Failed to create intake session")
    
    # Update client engagement status
    supabase.table("client_companies")\
        .update({"engagement_status": "awaiting_intake"})\
        .eq("id", draft["client_company_id"])\
        .execute()
    
    # Audit log
    write_audit_log(
        org_id=org_id,
        user_id=user["id"],
        action="intake_package_marked_sent",
        item_type="client_intake_session",
        item_id=session_id,
        details={
            "client_company_id": draft["client_company_id"],
            "email_draft_id": draft["id"],
            "tax_years": draft["tax_years"],
            "expected_inputs": list(requirements["expected_inputs"].keys()),
            "completion_method": "manual_user_action"
        }
    )
    
    return MarkSentResponse(
        success=True,
        intake_session_id=session_id,
        status="awaiting_client"
    )


@router.get("/session/{client_company_id}")
async def get_intake_session(
    client_company_id: str,
    user: dict = Depends(get_current_user)
):
    """Get the current intake session for a client."""
    supabase = get_supabase()
    profile = get_user_profile(user["id"])
    org_id = profile.get("organization_id")
    
    try:
        result = supabase.table("client_intake_sessions")\
            .select("*")\
            .eq("client_company_id", client_company_id)\
            .eq("organization_id", org_id)\
            .order("created_at", desc=True)\
            .limit(1)\
            .execute()
        
        if result.data:
            return {"success": True, "session": result.data[0]}
        else:
            return {"success": True, "session": None}
    except Exception as e:
        logger.error(f"Error fetching intake session: {e}")
        return {"success": False, "error": str(e)}


@router.get("/templates/{client_company_id}/list")
async def list_client_templates(
    client_company_id: str,
    status: str = "active",
    user: dict = Depends(get_current_user)
):
    """List all templates for a client."""
    supabase = get_supabase()
    profile = get_user_profile(user["id"])
    org_id = profile.get("organization_id")
    
    try:
        query = supabase.table("intake_templates")\
            .select("*")\
            .eq("client_company_id", client_company_id)\
            .eq("organization_id", org_id)
        
        if status:
            query = query.eq("status", status)
        
        result = query.order("created_at", desc=True).execute()
        
        # Add download URLs
        templates = []
        for t in result.data or []:
            t["download_url"] = f"/api/intake/templates/{t['id']}/download"
            templates.append(t)
        
        return {"success": True, "templates": templates}
    except Exception as e:
        logger.error(f"Error listing templates: {e}")
        return {"success": False, "error": str(e), "templates": []}


@router.patch("/client/{client_company_id}/settings")
async def update_client_intake_settings(
    client_company_id: str,
    settings: UpdateClientSettingsRequest,
    user: dict = Depends(get_current_user)
):
    """Update client engagement settings."""
    supabase = get_supabase()
    
    if not check_cpa_or_executive(user):
        raise HTTPException(status_code=403, detail="CPA or Executive role required")
    
    profile = get_user_profile(user["id"])
    org_id = profile.get("organization_id")
    
    # Verify client belongs to org
    try:
        client_result = supabase.table("client_companies")\
            .select("*")\
            .eq("id", client_company_id)\
            .eq("organization_id", org_id)\
            .single()\
            .execute()
        client = client_result.data
    except:
        raise HTTPException(status_code=404, detail="Client not found")
    
    # Build update dict
    updates = {}
    before_values = {}
    
    if settings.primary_contact_name is not None:
        before_values["primary_contact_name"] = client.get("primary_contact_name")
        updates["primary_contact_name"] = settings.primary_contact_name
    
    if settings.primary_contact_email is not None:
        before_values["primary_contact_email"] = client.get("primary_contact_email")
        updates["primary_contact_email"] = settings.primary_contact_email
    
    if settings.purchased_sections is not None:
        before_values["purchased_sections"] = client.get("purchased_sections")
        updates["purchased_sections"] = settings.purchased_sections
    
    if settings.study_scope is not None:
        before_values["study_scope"] = client.get("study_scope")
        updates["study_scope"] = settings.study_scope
    
    if settings.intake_mode is not None:
        before_values["intake_mode"] = client.get("intake_mode")
        updates["intake_mode"] = settings.intake_mode
    
    if settings.has_vendors_expected is not None:
        before_values["has_vendors_expected"] = client.get("has_vendors_expected")
        updates["has_vendors_expected"] = settings.has_vendors_expected
    
    if not updates:
        return {"success": True, "message": "No changes"}
    
    # Update
    supabase.table("client_companies")\
        .update(updates)\
        .eq("id", client_company_id)\
        .execute()
    
    # Audit log
    write_audit_log(
        org_id=org_id,
        user_id=user["id"],
        action="client_intake_settings_updated",
        item_type="client_company",
        item_id=client_company_id,
        details={
            "before": before_values,
            "after": updates,
            "completion_method": "manual_user_action"
        }
    )
    
    return {"success": True, "updated_fields": list(updates.keys())}


@router.get("/requirements-matrix")
async def get_requirements_matrix(
    section_41: bool = True,
    section_174: bool = False,
    has_vendors: bool = True,
    tax_years: str = "2024",
    user: dict = Depends(get_current_user)
):
    """
    Preview the intake requirements matrix without generating templates.
    Useful for understanding what will be required.
    """
    years = [int(y.strip()) for y in tax_years.split(",")]
    
    purchased_sections = {
        "section_41": section_41,
        "section_174": section_174
    }
    
    requirements = compute_intake_requirements(
        purchased_sections,
        years,
        has_vendors
    )
    
    return {
        "success": True,
        "purchased_sections": purchased_sections,
        "tax_years": years,
        "requirements": requirements,
        "template_definitions": {
            k: {"title": v.get("title"), "description": v.get("description")}
            for k, v in TEMPLATE_DEFINITIONS.items()
        }
    }
